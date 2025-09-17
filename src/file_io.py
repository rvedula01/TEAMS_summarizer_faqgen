# -*- coding: utf-8 -*-
"""
Created on Fri Jun  6 15:16:04 2025

@author: ShivakrishnaBoora
"""

import os
from typing import List
import docx
# import docx
from docx import Document
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

def iter_block_items(parent):
    """
    Yield each paragraph and table in document order.
    """
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def extract_images_from_paragraph(paragraph, output_dir='temp_images'):
    """Extract images from a paragraph by checking for graphic elements in XML."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    image_paths = []
    
    # Check if paragraph contains any images by looking for 'graphic' in XML
    if hasattr(paragraph, '_element') and hasattr(paragraph._element, 'xml'):
        if 'graphic' in paragraph._element.xml.lower():
            # Get all runs in the paragraph
            for run in paragraph.runs:
                if hasattr(run, '_element') and hasattr(run._element, 'xml'):
                    xml_str = run._element.xml
                    if 'graphic' in xml_str.lower():
                        try:
                            # Extract image data
                            import re
                            from docx.oxml.ns import qn
                            
                            # Find all embedded objects
                            for elem in run._element.xpath('.//*[contains(name(), "blip")]'):
                                embed_id = elem.get(qn('r:embed') or qn('r:link'))
                                if embed_id and embed_id in paragraph.part.related_parts:
                                    image_part = paragraph.part.related_parts[embed_id]
                                    image_ext = image_part.content_type.split('/')[-1]
                                    
                                    # Save image
                                    image_filename = f"image_{len(os.listdir(output_dir)) + 1}.{image_ext}"
                                    image_path = os.path.join(output_dir, image_filename)
                                    with open(image_path, 'wb') as f:
                                        f.write(image_part.blob)
                                    
                                    image_paths.append(image_path)
                        except Exception as e:
                            print(f"Warning: Could not extract image: {str(e)}")
                            continue
    
    return image_paths

def load_text_from_docx(path: str, encoding: str = 'utf-8') -> str:
    """
    Read a .docx and return its content with image placeholders:
    - Paragraphs become plain text lines.
    - Tables become Markdown tables (headers + separator + rows).
    - Images are saved and replaced with placeholders.
    - Consecutive lines from the same speaker are merged with the earliest timestamp.
    """
    try:
        doc = Document(path)
        segments = []
        last_speaker = None
        last_timestamp = None
        current_segment = []
        
        def add_segment():
            if current_segment and last_speaker and last_timestamp:
                merged_text = '; '.join(line for line in current_segment if line)
                segments.append(f"{last_timestamp} {last_speaker} {merged_text}")
        
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                # Check for images in the paragraph first
                image_paths = extract_images_from_paragraph(block)
                for img_path in image_paths:
                    segments.append(f"[IMAGE:{img_path}]")
                
                text = block.text.strip()
                if not text and not image_paths:
                    continue
                    
                # Check if line contains a timestamp and speaker (format: "HH:MM:SS AM/PM Speaker: ")
                timestamp_match = re.match(r'^(\d{1,2}:\d{2}(?::\d{2})? (?:AM|PM))\s+([^:]+):\s*(.*)', text, re.IGNORECASE)
                
                if timestamp_match:
                    # If we have a pending segment, add it before starting a new one
                    if current_segment and last_speaker:
                        add_segment()
                        current_segment = []
                    
                    timestamp, speaker, content = timestamp_match.groups()
                    last_timestamp = timestamp
                    last_speaker = speaker.strip()
                    if content.strip():
                        current_segment.append(content.strip())
                elif text:  # Only process non-empty text
                    # If it's a continuation line, add to current segment
                    if last_speaker and last_timestamp:
                        current_segment.append(text.strip())
                    else:
                        segments.append(text)
                        
            elif isinstance(block, Table):
                # If we have a pending segment, add it before processing the table
                if current_segment and last_speaker:
                    add_segment()
                    current_segment = []
                
                # Process table as before
                rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
                if not rows:
                    continue
                header = rows[0]
                separator = ['---'] * len(header)
                md = []
                md.append('| ' + ' | '.join(header) + ' |')
                md.append('| ' + ' | '.join(separator) + ' |')
                for data_row in rows[1:]:
                    md.append('| ' + ' | '.join(data_row) + ' |')
                segments.append('\n'.join(md))
        
        # Add the last segment if it exists
        if current_segment and last_speaker:
            add_segment()
        
        return '\n\n'.join(segments)
    except Exception as e:
        raise Exception(f"Error reading file {path}: {str(e)}")

def save_text_to_docx(text: str, output_path: str) -> None:
    """
    Save text to a .docx file with improved table handling.
    - Processes only the last table in the document
    - Maintains all text content before the table
    - Applies consistent formatting and styling
    - Handles edge cases for empty or malformed input
    """
    if not text.strip():
        raise ValueError("Cannot save empty content to document")

    # Clean the text and split into lines
    cleaned_text = re.sub(r"[`'\"""'']+", "", text)
    lines = [line.strip() for line in cleaned_text.splitlines() if line.strip()]
    
    # Find the last table in the document
    table_start = -1
    table_end = -1
    in_table = False
    
    for i, line in enumerate(lines):
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                table_start = i
                in_table = True
            table_end = i + 1
        elif in_table:
            in_table = False
    
    # Extract content sections
    content_before = lines[:table_start] if table_start >= 0 else lines
    table_content = lines[table_start:table_end] if table_start >= 0 else []
    
    # Create document
    doc = Document()
    
    # Add content before table
    if content_before:
        current_paragraph = None
        for line in content_before:
            if line:  # Only add non-empty lines
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph(line)
                else:
                    current_paragraph.add_run('\n' + line)
    
    # Add table if it exists
    if table_content:
        try:
            # Parse table rows
            rows = []
            for line in table_content:
                if line.startswith('|') and line.endswith('|'):
                    # Remove leading/trailing | and split by |
                    cells = [cell.strip() for cell in line[1:-1].split('|')]
                    rows.append(cells)
            
            if rows:  # Only proceed if we have valid rows
                # Skip separator row (the one with ---)
                rows = [row for row in rows if not all(cell.replace('-', '').strip() == '' for cell in row)]
                
                if rows:  # If we still have rows after filtering
                    # Create table with proper dimensions
                    num_cols = max(len(row) for row in rows)
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Table Grid'  # Add grid lines
                    
                    # Fill table cells
                    for i, row in enumerate(rows):
                        for j, cell_text in enumerate(row):
                            if j < num_cols:  # Ensure we don't exceed column count
                                cell = table.cell(i, j)
                                cell.text = cell_text
                                
                                # Apply cell formatting
                                for paragraph in cell.paragraphs:
                                    paragraph.style = 'Table Text'
        except Exception as e:
            # If table processing fails, add the raw table content as text
            doc.add_paragraph("\n".join(table_content))
    
    # Ensure output directory exists and save
    os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
    doc.save(output_path)

def ensure_file_exists(path: str) -> None:
    """
    Raise FileNotFoundError if `path` does not exist.
    """
    if not os.path.isfile(path):
        raise FileNotFoundError(f"File not found: {path}")
