# -*- coding: utf-8 -*-
"""
Created on Fri Jun  6 15:18:14 2025

@author: ShivakrishnaBoora
"""

import os
import sys
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables from .env file in the project root
dotenv_path = Path(__file__).parent.parent / '.env'
load_dotenv(dotenv_path=dotenv_path)

import re
import json
import ast
import platform
import threading
from datetime import datetime, timedelta
import base64
import tempfile
import nltk

# Download NLTK data if not already downloaded
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab', quiet=True)

# Windows-specific imports
if platform.system() == 'Windows':
    import pythoncom
    import win32com.client


import streamlit as st
import glob
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import numpy as np
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from file_io import (
    load_text_from_docx,
    save_text_to_docx,
    ensure_file_exists,
    extract_images_from_paragraph,
    iter_block_items
)
from text_processing import clean_transcript, summarize_transcript, chunked_clean_and_summarize
from text_processing import chunked_clean_and_summarize, _split_raw_into_chunks
from aggregator import read_chunk_summaries, summarize_timeline, summarize_table, extract_single_markdown_table
from openai_client import call_openai_chat

# Import the merge functionality
from merge_chat_transcript import merge_chat_and_transcript, convert_chat_only

import pandas as pd
from typing import List, Dict, Any
# Workaround for PyTorch and Streamlit file watcher conflict
import os
os.environ['STREAMLIT_SERVER_FILE_WATCHER_TYPE'] = 'none'

os.environ['STREAMLIT_FILE_WATCHER'] = 'none'

#st.write("API Key exists:", os.getenv("OPENAI_API_KEY") is not None)

def replace_png_placeholder_fixed(match):
    """
    Fixed PNG image placeholder replacement with better path extraction.
    """
    img_path = match.group(1).strip()
    
    # Clean up the path by removing any trailing duplicate image references
    # Handle cases like: "temp_images\image_76.png [Image:Temp_Images\Image_76.Png"
    if ' [Image:' in img_path or ' [IMAGE:' in img_path:
        # Split at the first occurrence of ' [Image:' or ' [IMAGE:'
        img_path = re.split(r'\s+\[(?:Image|IMAGE):', img_path, flags=re.IGNORECASE)[0]
    
    img_path = img_path.strip()
    
    print(f"DEBUG: Cleaned PNG image path: '{img_path}'")
    
    if not img_path:
        return "[Invalid image path]"
    
    # Generate path variations to try
    path_variations = [
        img_path,
        img_path.replace('\\', '/'),
        os.path.normpath(img_path),
        os.path.abspath(img_path)
    ]
    
    # If path doesn't end with .png, add it
    for i, path in enumerate(path_variations):
        if not path.lower().endswith('.png'):
            path_variations.append(path + '.png')
    
    print(f"DEBUG: Trying cleaned paths: {path_variations}")
    
    for path in path_variations:
        try:
            if os.path.exists(path):
                print(f"DEBUG: Found PNG at: {path}")
                # Read and encode PNG as base64
                with open(path, "rb") as img_file:
                    img_data = img_file.read()
                    img_base64 = base64.b64encode(img_data).decode()
                    
                    # Return base64 encoded PNG for HTML display
                    return f'<img src="data:image/png;base64,{img_base64}" style="max-width: 100%; max-height: 400px; display: block; margin: 10px auto; border: 1px solid #ddd;" alt="PNG from transcript">'
        except Exception as e:
            print(f"DEBUG: Error processing PNG {path}: {e}")
            continue
    
    # If no PNG found, return placeholder text
    print(f"DEBUG: PNG not found for cleaned path: '{img_path}'")
    return f"[PNG not found: {os.path.basename(img_path)}]"

def process_first_image_only_for_display(text_content, for_word_doc=False):
    """
    Process only the first image for display and remove all other image placeholders.
    This function should be used wherever you're displaying the final summary.
    
    Args:
        text_content: The text content to process
        for_word_doc: If True, removes HTML tags; if False, keeps HTML for web display
    """
    import re
    
    # Track if an image has been processed
    image_processed = False
    processed_lines = []
    
    # PNG image pattern
    png_pattern = r'\[(?:IMAGE|Image):\s*([^]]+?)(?:\s*\[Image:[^]]+\])?\]'
    
    # Split text by lines to preserve structure
    lines = text_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            processed_lines.append('')
            continue
        
        # Check if line contains PNG image references
        if re.search(png_pattern, line, re.IGNORECASE):
            if not image_processed:
                # Process only the first image
                image_processed = True
                
                # Process the first image in the line
                processed_line = re.sub(
                    png_pattern, 
                    replace_png_placeholder_fixed,
                    line, 
                    count=1,  # Only replace first occurrence
                    flags=re.IGNORECASE
                )
                
                # Remove any remaining image placeholders from the same line
                processed_line = re.sub(png_pattern, '', processed_line, flags=re.IGNORECASE)
                
                # If this is for Word document, remove HTML tags
                if for_word_doc:
                    processed_line = re.sub(r'<img[^>]*>', '[Image embedded above]', processed_line)
                    processed_line = re.sub(r'<[^>]+>', '', processed_line)
                
                processed_lines.append(processed_line)
            else:
                # Remove all image placeholders from subsequent lines
                clean_line = re.sub(png_pattern, '', line, flags=re.IGNORECASE).strip()
                if clean_line:
                    processed_lines.append(clean_line)
        else:
            # No images in this line, add as-is
            processed_lines.append(line)
    
    return '\n'.join(processed_lines)

def _is_meaningful_text(text: str, min_words: int = 5) -> bool:
    """Check if text is meaningful (contains more than min_words words)."""
    if not text or not text.strip():
        return False
    words = text.split()
    return len(words) > min_words

def _is_mostly_code_or_gibberish(text: str) -> bool:
    """Check if text is mostly code or gibberish."""
    if not text:
        return True
    
    # Count non-alphabetic characters
    alpha_chars = sum(c.isalpha() for c in text)
    total_chars = max(1, len(text) - text.count(' '))  # Exclude spaces
    alpha_ratio = alpha_chars / total_chars
    
    # If less than 30% alphabetic, likely code or gibberish
    return alpha_ratio < 0.3

def _split_text_into_chunks(text: str, max_chunk_size: int = 4000) -> List[str]:
    """
    Split text into meaningful chunks, trying to keep related content together.
    
    Args:
        text: The text to split into chunks
        max_chunk_size: Maximum size of each chunk in characters (default: 4000)
        
    Returns:
        List of meaningful text chunks, each <= max_chunk_size
    """
    if not text or not text.strip():
        return []
    
    # First, split by major sections (double newlines)
    sections = [s.strip() for s in text.split('\n\n') if s.strip()]
    chunks = []
    current_chunk = []
    current_size = 0
    
    for section in sections:
        if not _is_meaningful_text(section) or _is_mostly_code_or_gibberish(section):
            continue
            
        section_size = len(section)
        
        # If section is too large, try to split by sentences
        if section_size > max_chunk_size:
            sentences = nltk.sent_tokenize(section)
            current_sent_chunk = []
            current_sent_size = 0
            
            for sent in sentences:
                sent = sent.strip()
                if not _is_meaningful_text(sent):
                    continue
                    
                sent_size = len(sent)
                
                # If adding this sentence would exceed the chunk size, finalize current chunk
                if current_sent_size + sent_size > max_chunk_size and current_sent_chunk:
                    chunk_text = ' '.join(current_sent_chunk)
                    if _is_meaningful_text(chunk_text):
                        chunks.append(chunk_text)
                    current_sent_chunk = []
                    current_sent_size = 0
                    
                current_sent_chunk.append(sent)
                current_sent_size += sent_size + 1
            
            # Add remaining sentences if any
            if current_sent_chunk:
                chunk_text = ' '.join(current_sent_chunk)
                if _is_meaningful_text(chunk_text):
                    chunks.append(chunk_text)
        
        # If section fits in current chunk, add it
        elif current_size + section_size <= max_chunk_size:
            current_chunk.append(section)
            current_size += section_size + 2  # +2 for newlines
        
        # Otherwise, finalize current chunk and start a new one
        else:
            if current_chunk:
                chunk_text = '\n\n'.join(current_chunk)
                if _is_meaningful_text(chunk_text):
                    chunks.append(chunk_text)
            current_chunk = [section]
            current_size = section_size
    
    # Add the last chunk if not empty
    if current_chunk:
        chunk_text = '\n\n'.join(current_chunk)
        if _is_meaningful_text(chunk_text):
            chunks.append(chunk_text)
    
    # Log chunk statistics
    print(f"\n=== Chunking Statistics ===")
    print(f"Original text length: {len(text)} characters")
    print(f"Number of chunks: {len(chunks)}")
    print(f"Average chunk size: {sum(len(c) for c in chunks) / max(1, len(chunks)):.0f} characters")
    
    if chunks:
        print("\nSample chunks:")
        for i, chunk in enumerate(chunks[:3], 1):
            preview = chunk[:100].replace('\n', ' ')
            print(f"Chunk {i} ({len(chunk)} chars): {preview}...")
        if len(chunks) > 3:
            print(f"... and {len(chunks) - 3} more chunks")
    
    return chunks

# FAQ Extraction Functions
def extract_faqs(text: str, max_chunk_size: int = 4000) -> List[Dict[str, str]]:
    """
    Extract FAQs (questions and answers) from the given text using OpenAI with chunking.
    Uses a more efficient chunking strategy to reduce the number of API calls.
    
    Args:
        text: The text to extract FAQs from
        max_chunk_size: Maximum size of each chunk in characters (default: 4000)
        
    Returns:
        List of dictionaries containing questions and their corresponding answers
    """
    if not text or not text.strip():
        print("Warning: Empty text provided to extract_faqs")
        return []
    
    try:
        print("\n=== Starting FAQ Extraction ===")
        print(f"Input text length: {len(text)} characters")
        
        # Split text into meaningful chunks
        chunks = _split_text_into_chunks(text, max_chunk_size)
        
        if not chunks:
            print("Warning: No meaningful chunks could be created from the text")
            return []
        
        all_faqs = []
        total_chunks = len(chunks)
        
        print(f"\nProcessing {total_chunks} chunks...")
        
        for idx, chunk in enumerate(chunks, 1):
            try:
                print(f"\n=== Chunk {idx}/{total_chunks} ===")
                print(f"Size: {len(chunk):,} characters")
                
                # Process the chunk
                chunk_faqs = _process_chunk_for_faqs(chunk)
                
                if chunk_faqs:
                    print(f"Found {len(chunk_faqs)} FAQs in this chunk")
                    all_faqs.extend(chunk_faqs)
                else:
                    print("No FAQs found in this chunk")
                
            except Exception as chunk_error:
                print(f"Error processing chunk {idx}: {str(chunk_error)}")
                import traceback
                traceback.print_exc()
                continue
        
        # Remove duplicates while preserving order
        unique_faqs = []
        seen_questions = set()
        
        for faq in all_faqs:
            if not isinstance(faq, dict) or 'question' not in faq:
                continue
                
            question = faq['question'].strip().lower()
            if question and question not in seen_questions:
                seen_questions.add(question)
                unique_faqs.append(faq)
        
        # Sort FAQs by the order they first appeared in the text
        def get_first_occurrence(faq):
            question = faq['question'].lower()
            return text.lower().find(question)
            
        unique_faqs.sort(key=get_first_occurrence)
        
        print(f"\n=== FAQ Extraction Complete ===")
        print(f"Total unique FAQs found: {len(unique_faqs)}")
        
        if unique_faqs:
            print("\nSample FAQs:")
            for i, faq in enumerate(unique_faqs[:3], 1):
                print(f"{i}. Q: {faq.get('question', '')}")
                print(f"   A: {faq.get('answer', '')[:100]}...")
            if len(unique_faqs) > 3:
                print(f"... and {len(unique_faqs) - 3} more")
        
        return unique_faqs
        
    except Exception as e:
        print(f"\n!!! CRITICAL ERROR in extract_faqs: {str(e)}")
        import traceback
        traceback.print_exc()
        return []
        raise

def _process_chunk_for_faqs(chunk: str) -> List[Dict[str, str]]:
    """
    Helper function to process a single chunk of text for FAQs using OpenAI.
    """
    if not chunk or not chunk.strip():
        print("Warning: Empty chunk provided to _process_chunk_for_faqs")
        return []
        
    try:
        print("\n=== Starting FAQ Extraction ===")
        print(f"Chunk size: {len(chunk)} characters")
        
        # Safety check - if chunk is still too large, truncate it with a warning
        max_allowed = 3000  # Conservative limit
        if len(chunk) > max_allowed:
            print(f"Warning: Chunk size ({len(chunk)}) exceeds maximum allowed size ({max_allowed}). Truncating...")
            chunk = chunk[:max_allowed] + "\n[Content truncated due to length]"
            
        print(f"Processing chunk (first 100 chars): {chunk[:100]}...")
        
        # First, try to extract any existing Q&A pairs directly from the text
        qa_pairs = _extract_existing_qa_pairs(chunk)
        if qa_pairs:
            print(f"Found {len(qa_pairs)} Q&A pairs directly in the text")
            return qa_pairs
        
        print("No direct Q&A pairs found, using LLM to generate them...")
            
        # If no direct Q&A pairs found, use the LLM to generate them
        prompt = f"""
You are a technical incident data extraction assistant. Your task is to extract high-quality, technical questions and their corresponding answers from incident call transcripts or technical discussions.

## INCLUSION CRITERIA:
- Technical questions about systems, services, or processes
- Incident-related queries about root causes, impacts, or resolutions
- Questions about system status, configurations, or behaviors
- Technical troubleshooting steps and their outcomes
- Questions about error messages or system logs
- Technical decisions and their rationales

## EXCLUSION CRITERIA (DO NOT INCLUDE):
- Greetings, casual conversation, or social pleasantries
- Administrative or logistical discussions
- Simple confirmations (e.g., "Is the system up?" - "Yes")
- Incomplete or ambiguous questions without clear technical answers
- Questions where the answer is just "yes", "no", or "I don't know"
- Questions about meeting logistics or scheduling
- Questions where the answer is just "I'm checking" or similar non-answers
- Questions that are actually statements or requests
- Questions about non-technical topics

## QUALITY REQUIREMENTS:
1. Questions must be complete and technically specific
2. Answers must provide substantial technical information or context
3. Both question and answer should be clear and understandable on their own
4. Remove any filler words or phrases from both questions and answers
5. If a question is technical but the answer is non-substantive (e.g., "I'll check"), exclude it
6. For yes/no questions, only include if the answer provides detailed technical explanation

## EXAMPLES OF WHAT TO EXCLUDE:
- Q: "Is the database up?" 
  A: "Yes"
  
- Q: "Can you check the logs?"
  A: "I'm looking at them now"
  
- Q: "Are we ready to start?"
  A: "Yes, let's begin"

## EXAMPLES OF WHAT TO INCLUDE:
- Q: "What was the root cause of the database outage?"
  A: "The primary database server ran out of disk space due to unrotated log files. We've added monitoring and automated log rotation to prevent recurrence."

- Q: "Which services were affected by the network partition?"
  A: "The authentication service and payment processing were impacted between 14:30 and 15:45 UTC. We've implemented circuit breakers to limit blast radius in future events."

## OUTPUT FORMAT:
Return a valid JSON array where each element is an object with two fields:
- "question": the technical question (exact wording or appropriately formulated)
- "answer": the complete, relevant answer with sufficient technical context

Output only the JSON. Do not include any other explanation or text.

Text to analyze:
{chunk}
"""
        
        print("Sending request to OpenAI...")
        response = call_openai_chat(prompt=prompt, model="gpt-4o")
        
        if not response:
            print("ERROR: Empty response from OpenAI")
            return []
            
        # Print raw response for debugging
        print(f"\n=== Raw response from OpenAI ===")
        print(f"Response type: {type(response)}")
        print(f"Response length: {len(response)} characters")
        print(f"Response preview: {response[:200]}...")
        
        # Remove any code block markers and trim whitespace
        cleaned_content = response.replace("```json", "").replace("```", "").strip()
        
        # Try multiple parsing methods
        print("\nAttempting to parse response...")
        
        # First, try to find JSON array in the response
        json_start = cleaned_content.find('[')
        json_end = cleaned_content.rfind(']') + 1
        
        if json_start >= 0 and json_end > json_start:
            json_str = cleaned_content[json_start:json_end]
            print(f"Extracted JSON string (length: {len(json_str)}): {json_str[:100]}...")
            
            try:
                # Try json.loads first
                faqs = json.loads(json_str)
                print(f"Successfully parsed JSON with {len(faqs) if isinstance(faqs, list) else 0} items")
                
                if isinstance(faqs, list) and len(faqs) > 0:
                    # Validate each item in the list
                    valid_faqs = []
                    for i, item in enumerate(faqs):
                        if not isinstance(item, dict):
                            print(f"Warning: Item {i} is not a dictionary: {item}")
                            continue
                            
                        if 'question' not in item or 'answer' not in item:
                            print(f"Warning: Item {i} is missing required fields: {item}")
                            continue
                            
                        # Convert values to strings and clean them
                        question = str(item.get('question', '')).strip()
                        answer = str(item.get('answer', '')).strip()
                        
                        if question and answer:  # Only add if both fields are non-empty
                            valid_faqs.append({
                                'question': question,
                                'answer': answer
                            })
                    
                    if valid_faqs:
                        print(f"Successfully extracted {len(valid_faqs)} valid FAQs")
                        return valid_faqs
                    
                    print("No valid FAQs found after validation")
                else:
                    print("No FAQs found in the parsed JSON")
                    
            except json.JSONDecodeError as e:
                print(f"JSON decode error: {e}")
                # Fall through to try other parsing methods
            except Exception as e:
                print(f"Unexpected error parsing JSON: {e}")
                # Fall through to try other parsing methods
        
        # If we get here, try ast.literal_eval as a fallback
        print("Trying ast.literal_eval as fallback...")
        try:
            faqs = ast.literal_eval(cleaned_content)
            print(f"ast.literal_eval parsed content with {len(faqs) if isinstance(faqs, list) else 0} items")
            
            if isinstance(faqs, list) and len(faqs) > 0:
                # Validate each item in the list
                valid_faqs = []
                for i, item in enumerate(faqs):
                    if not isinstance(item, dict):
                        continue
                        
                    question = str(item.get('question', '')).strip()
                    answer = str(item.get('answer', '')).strip()
                    
                    if question and answer:  # Only add if both fields are non-empty
                        valid_faqs.append({
                            'question': question,
                            'answer': answer
                        })
                
                if valid_faqs:
                    print(f"Successfully extracted {len(valid_faqs)} valid FAQs using ast.literal_eval")
                    return valid_faqs
                
                print("No valid FAQs found after ast.literal_eval validation")
            else:
                print("No FAQs found in the ast.literal_eval result")
                
        except Exception as e:
            print(f"ast.literal_eval failed: {e}")
                
        # If we get here, all parsing attempts have failed
        print("\n=== FAQ Extraction Failed ===")
        print("All parsing attempts failed. Raw cleaned content:")
        print(cleaned_content[:500] + ("..." if len(cleaned_content) > 500 else ""))
        return []
            
    except Exception as e:
        print(f"Error processing chunk for FAQs: {e}")
        import traceback
        traceback.print_exc()
        return []

def _extract_existing_qa_pairs(text: str) -> List[Dict[str, str]]:
    """
    Helper function to extract existing Q&A pairs from text using pattern matching.
    """
    import re
    
    # Pattern to match Q: and A: or similar patterns
    qa_pattern = r'(?i)(?:Q|Question|Q\s*[:\.])\s*(.*?)(?=\n\s*(?:A|Answer|A\s*[:.])\s*|\Z)(?:\n\s*(?:A|Answer|A\s*[:.])\s*(.*?))(?=\n\s*(?:Q|Question|Q\s*[:.])|\Z)'
    
    matches = re.findall(qa_pattern, text, re.DOTALL)
    
    qa_pairs = []
    for match in matches:
        question = match[0].strip()
        answer = match[1].strip() if len(match) > 1 else ""
        if question and answer:
            qa_pairs.append({
                'question': question,
                'answer': answer
            })
    
    return qa_pairs

def create_faq_document(faqs: List[Dict[str, str]], output_path: str = "faq_document.docx") -> str:
    """
    Create a Word document with extracted FAQs.
    
    Args:
        faqs: List of FAQ dictionaries with 'question' and 'answer' keys
        output_path: Path for the output document
        
    Returns:
        Path to the created document
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Frequently Asked Questions (FAQs)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y')}").italic = True
    
    doc.add_paragraph()
    
    if faqs:
        for i, faq in enumerate(faqs, 1):
            # Add question
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(f"Q{i}: {faq.get('question', 'No question provided')}")
            question_run.bold = True
            question_run.font.size = Pt(12)
            
            # Add answer
            answer_para = doc.add_paragraph()
            answer_run = answer_para.add_run(f"A{i}: {faq.get('answer', 'No answer provided')}")
            answer_run.font.size = Pt(11)
            
            # Add spacing between Q&A pairs
            doc.add_paragraph()
    else:
        doc.add_paragraph("No FAQs were extracted from the provided text.")
    
    # Save the document
    doc.save(output_path)
    return output_path

# Function to extract action items from summary
def chunk_text(text, max_chunk_size=3500):
    """Split text into chunks of approximately max_chunk_size characters, trying to break at sentence boundaries."""
    if not text:
        return []
        
    # Split into sentences first
    sentences = text.split('. ')
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # If adding this sentence would exceed the chunk size, start a new chunk
        if current_chunk and len(current_chunk) + len(sentence) > max_chunk_size:
            chunks.append(current_chunk)
            current_chunk = sentence + ". "
        else:
            current_chunk += sentence + ". "
    
    # Add the last chunk if it's not empty
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

def extract_action_items(summary, max_chunk_size=3000):
    """
    Extract action items from the summary text using OpenAI with chunking.
    
    Args:
        summary: The summary text to extract action items from
        max_chunk_size: Maximum size of each chunk in characters
        
    Returns:
        List of dictionaries containing action items with responsible, action, and status
    """
    if not summary:
        return []
    
    # Split the summary into manageable chunks
    chunks = chunk_text(summary, max_chunk_size)
    all_action_items = []
    
    for i, chunk in enumerate(chunks, 1):
        try:
            print(f"Processing action items chunk {i}/{len(chunks)}")
            chunk_action_items = _process_chunk_for_action_items(chunk)
            all_action_items.extend(chunk_action_items)
        except Exception as e:
            print(f"Error processing action items chunk {i}: {str(e)}")
            continue
    
    # Remove duplicates (if any)
    unique_action_items = []
    seen = set()
    for item in all_action_items:
        # Create a unique key for each action item
        key = (item.get('action', '').lower().strip(), 
               item.get('responsible', '').lower().strip())
        if key not in seen:
            seen.add(key)
            unique_action_items.append(item)
    
    print(f"Extracted {len(unique_action_items)} unique action items from {len(chunks)} chunks")
    return unique_action_items

def _process_chunk_for_action_items(chunk):
    """Helper function to process a single chunk of text for action items"""
    from openai_client import call_openai_chat
    
    prompt = f"""Extract action items from the following meeting summary chunk.
    Return a JSON array of action items with this exact structure:
    
    [
        {{
            "responsible": "Team Name or 'Unassigned'",
            "action": "Detailed action description",
            "status": "Open/In Progress/Completed (default: 'Open')"
        }}
    ]
    
    Rules:
    1. Extract all action items, tasks, or follow-ups
    2. Keep the original wording but make actions clear and specific
    3. Include any mentioned deadlines or timeframes in the action text
    4. For 'responsible' field, ONLY use team names (e.g., 'Application Team', 'Security Team', 'DevOps Team')
       - NEVER use individual names like 'John', 'Sarah', 'Aneesh L 1.5 team' etc.
       - If no team is mentioned, use 'Unassigned'
    5. Status should be one of: 'Open', 'In Progress', or 'Completed'
    6. Return ONLY valid JSON, no other text or explanation
    
    Examples of good team names:
    - 'ABC Team' (not 'Neerav')
    - 'GBS Team' (not 'Praveen')
    
    Meeting Summary Chunk:
    {chunk}
    """
    
    response = call_openai_chat(prompt=prompt, model="gpt-4o")
    response = response.replace("```json", "").replace("```", "").strip()
    
    import json
    action_items = json.loads(response)
    
    # Ensure all required fields are present
    for item in action_items:
        item.setdefault('responsible', 'Unassigned')
        item.setdefault('status', 'Open')
        # Remove timestamp if it exists
        if 'timestamp' in item:
            del item['timestamp']
        
    return action_items

def extract_key_observations(summary, max_chunk_size=3000):
    """
    Extract key observations from the summary text using OpenAI with chunking.
    
    Args:
        summary: The summary text to extract observations from
        max_chunk_size: Maximum size of each chunk in characters
        
    Returns:
        List of observation strings
    """
    if not summary:
        return []
    
    # Split the summary into manageable chunks
    chunks = chunk_text(summary, max_chunk_size)
    all_observations = []
    
    for i, chunk in enumerate(chunks, 1):
        try:
            print(f"Processing observations chunk {i}/{len(chunks)}")
            chunk_observations = _process_chunk_for_observations(chunk)
            all_observations.extend(chunk_observations)
        except Exception as e:
            print(f"Error processing observations chunk {i}: {str(e)}")
            continue
    
    # Remove duplicates (if any)
    unique_observations = []
    seen = set()
    for obs in all_observations:
        # Normalize the observation for comparison
        normalized = ' '.join(obs.lower().split())
        if normalized not in seen:
            seen.add(normalized)
            unique_observations.append(obs)
    
    print(f"Extracted {len(unique_observations)} unique observations from {len(chunks)} chunks")
    return unique_observations

def _process_chunk_for_observations(chunk):
    """Helper function to process a single chunk of text for observations"""
    from openai_client import call_openai_chat
    
    prompt = """Extract key observations, decisions, or important points from the following meeting summary chunk.
    Pay special attention to any mentions of P1, P2, or tickets as these are high priority.
    
    Return a JSON array of observation strings with this exact structure:
    
    [
        "Observation 1 text",
        "Observation 2 text",
        ...
    ]
    
    Rules:
    1. Focus on important decisions, key points, and notable information
    2. Keep observations concise but meaningful
    3. Include any important context or reasoning
    4. Start each observation with a capital letter and end with a period
    5. Highlight any P1, P2, or ticket-related observations
    6. Return ONLY valid JSON, no other text or explanation
    
    Meeting Summary Chunk:
    {chunk}
    """.format(chunk=chunk)
    
    response = call_openai_chat(prompt=prompt, model="gpt-4o")
    response = response.replace("```json", "").replace("```", "").strip()
    
    import json
    observations = json.loads(response)
    
    # Ensure observations are properly formatted
    formatted_observations = []
    for obs in observations:
        if not isinstance(obs, str):
            obs = str(obs)
        # Ensure proper formatting
        obs = obs.strip()
        if not obs.endswith(('.', '!', '?')):
            obs += '.'
        obs = obs[0].upper() + obs[1:]  # Capitalize first letter
        formatted_observations.append(obs)
        
    return formatted_observations

def filter_priority_observations(observations: List[str]) -> List[str]:
    """
    Filter observations to only include those containing P1, P2, or ticket keywords.
    
    Args:
        observations: List of observation strings
        
    Returns:
        Filtered list of observations containing priority keywords
    """
    if not observations:
        return []
        
    keywords = ['P1', 'P2', 'ticket', 'Ticket']
    return [obs for obs in observations if any(keyword in obs for keyword in keywords)]

# Thread-local storage for COM objects
class COMWrapper:
    def __init__(self):
        self.initialized = False
        self.is_windows = platform.system() == 'Windows'
        
    def __enter__(self):
        # Only initialize COM on Windows
        if self.is_windows:
            try:
                pythoncom.CoInitialize()
                self.initialized = True
            except Exception as e:
                print(f"Warning: Could not initialize COM: {e}")
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        # Only uninitialize if we were the ones who initialized it
        if self.initialized and self.is_windows:
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            self.initialized = False

def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Convert a Word document to PDF.
    On Windows, uses Microsoft Word for better formatting.
    On other platforms, creates a simple PDF with the text content.
    """
    try:
        # Create the output directory if it doesn't exist
        os.makedirs(os.path.dirname(os.path.abspath(pdf_path)), exist_ok=True)
        
        if platform.system() == 'Windows':
            # Use Word on Windows if available
            try:
                with COMWrapper() as wrapper:
                    word = win32com.client.DispatchEx('Word.Application')
                    doc = word.Documents.Open(os.path.abspath(docx_path))
                    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 is the code for PDF format
                    doc.Close()
                    word.Quit()
                
                # Verify the PDF was created
                if os.path.exists(pdf_path):
                    return
                
            except Exception as e:
                print(f"Warning: Could not use Word for PDF conversion: {e}")
                # Fall through to the cross-platform method
        
        # Cross-platform fallback: Create a simple PDF with the text content
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.units import inch
        from docx import Document
        
        # Read the DOCX file
        doc = Document(docx_path)
        
        # Create PDF
        doc_pdf = SimpleDocTemplate(pdf_path, pagesize=letter,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=72)
        
        # Container for the 'Flowable' objects
        story = []
        styles = getSampleStyleSheet()
        
        # Add each paragraph from the docx to the PDF
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                # Convert docx paragraph to a PDF paragraph
                p = Paragraph(para.text, styles['Normal'])
                story.append(p)
        
        # Add each table from the docx to the PDF
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            p = Paragraph(para.text, styles['Normal'])
                            story.append(p)
        
        # Build the PDF
        doc_pdf.build(story)
        
    except Exception as e:
        print(f"Error in convert_docx_to_pdf: {e}")
        # Create a minimal PDF with an error message
        try:
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.drawString(100, 750, "Document Conversion Notice")
            c.drawString(100, 730, "The original document could not be converted with formatting.")
            c.drawString(100, 710, "Please check the original document for proper formatting.")
            c.save()
        except Exception as fallback_error:
            print(f"Failed to create fallback PDF: {fallback_error}")
            raise

def extract_meeting_datetime(text):
    # Match the specific format: Month DD, YYYY, HH:MMAM
    date_pattern = r'^(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{1,2},\s+\d{4},\s+(\d{1,2}):(\d{2})(AM|PM)'
    match = re.search(date_pattern, text, re.IGNORECASE)
    if match:
        # Extract just the date part (without time)
        date_str = text.split(',')[0] + ', ' + text.split(',')[1]
        # Get hour, minute, and AM/PM separately
        hour = int(match.group(1))
        minute = int(match.group(2))
        ampm = match.group(3).upper()
        
        # Convert to 24-hour format
        if ampm == 'PM' and hour != 12:
            hour = (hour % 12) + 12
        elif ampm == 'AM' and hour == 12:
            hour = 0
        
        # Validate time
        if hour < 0 or hour > 23 or minute < 0 or minute > 59:
            return None, None
            
        # Return date without time suffix and time in 24-hour format
        return date_str, f"{hour:02d}:{minute:02d}"
    return None, None


def format_timestamp(timestamp_str, meeting_date, meeting_start_time):
    """
    Convert MM:SS or H:MM:SS timestamp to full date format: Month DD, YYYY, HH:MM:SS AM/PM

    Args:
        timestamp_str: MM:SS or H:MM:SS/HH:MM:SS format timestamp string
        meeting_date: Date string in format "Month DD, YYYY"
        meeting_start_time: Start time in 24-hour format "HH:MM"

    Returns:
        Formatted datetime string or original timestamp if invalid
    """
    timestamp_str = str(timestamp_str).strip('`\'" ')

    try:
        # Parse meeting start time
        start_time = datetime.strptime(meeting_start_time, "%H:%M")

        # Determine if it's MM:SS or H:MM:SS/HH:MM:SS
        if re.match(r'^\d{1,2}:\d{2}$', timestamp_str):  # MM:SS
            minutes, seconds = map(int, timestamp_str.split(':'))
            delta = timedelta(minutes=minutes, seconds=seconds)
        elif re.match(r'^\d{1,2}:\d{2}:\d{2}$', timestamp_str):  # H:MM:SS or HH:MM:SS
            hours, minutes, seconds = map(int, timestamp_str.split(':'))
            delta = timedelta(hours=hours, minutes=minutes, seconds=seconds)
        else:
            print(f"Invalid timestamp format: '{timestamp_str}'")
            return timestamp_str

        # Build meeting datetime (date + start time)
        base_datetime = datetime.strptime(meeting_date, "%B %d, %Y")
        meeting_datetime = datetime.combine(base_datetime.date(), start_time.time()) + delta

        # Return formatted output
        return meeting_datetime.strftime("%B %d, %Y, %I:%M:%S %p")

    except Exception as e:
        print(f"Error formatting timestamp '{timestamp_str}': {e}")
        return timestamp_str


def ensure_newlines(text):
    """
    Insert a newline before each timestamp (MM:SS or H:MM:SS or HH:MM:SS)
    """
    return re.sub(r'(?<!^)(?<!\n)(\d{1,2}:\d{2}(?::\d{2})?)', r'\n\1', text)


def reformat_timestamp(timestamp_str):
    """
    Convert 'Month DD, YYYY, HH:MM:SS AM/PM' to 'MM/DD/YYYY, HH:MM:SS' (24-hour format).
    Passes through H:MM:SS-style times as-is.

    Args:
        timestamp_str: Full timestamp or short time format

    Returns:
        Reformatted timestamp or original string if parsing fails
    """
    try:
        # Pass through time-only formats (H:MM:SS or HH:MM:SS)
        if re.match(r'^\d{1,2}:\d{2}:\d{2}$', timestamp_str):
            return timestamp_str

        # Try full timestamp parsing
        try:
            dt = datetime.strptime(timestamp_str, "%B %d, %Y, %I:%M:%S %p")
        except ValueError:
            dt = datetime.strptime(timestamp_str, "%B %d, %Y %I:%M:%S %p")

        return dt.strftime("%m/%d/%Y, %H:%M:%S")

    except Exception:
        return timestamp_str

# Download NLTK punkt tokenizer if not already downloaded
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

SOURCE_TRANSCRIPT_PATH = "INC0671705 -- Transcript.docx"
MAX_CHARS_PER_CHUNK = 3000

def count_pages_in_docx(docx_path):
    """
    Count the number of pages in a Word document by detecting page breaks.
    
    This function attempts to count pages by looking for explicit page breaks
    in the document, which provides a more accurate count than estimation.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import zipfile
        import os
        
        # Default to 1 page if anything goes wrong
        default_pages = 1
        
        try:
            # First try to extract page count from document properties if available
            try:
                with zipfile.ZipFile(docx_path) as z:
                    # Look for the app.xml file which might contain page count
                    if 'docProps/app.xml' in z.namelist():
                        with z.open('docProps/app.xml') as f:
                            import xml.etree.ElementTree as ET
                            tree = ET.parse(f)
                            root = tree.getroot()
                            # Look for Pages element in app.xml
                            pages_elem = root.find('{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Pages')
                            if pages_elem is not None and pages_elem.text.isdigit():
                                return int(pages_elem.text)
            except Exception as e:
                print(f"Could not extract page count from document properties: {e}")
            
            # If we couldn't get page count from properties, try counting page breaks
            doc = Document(docx_path)
            page_count = 1  # Start with 1 page (the first page doesn't need a page break)
            
            # Look for explicit page breaks in paragraphs
            for para in doc.paragraphs:
                # Check for page break in paragraph properties
                pPr = para._p.get_or_add_pPr()
                if pPr.pageBreakBefore is not None:
                    page_count += 1
                
                # Check for page breaks in runs
                for run in para.runs:
                    if run._element is not None and run._element.get_or_add_rPr().get_or_add_lastRenderedPageBreak() is not None:
                        page_count += 1
            
            # Also check section breaks which might indicate page breaks
            for section in doc.sections:
                # If section starts on a new page, increment page count
                if hasattr(section, 'start_type') and str(section.start_type) == 'NEW_PAGE':
                    page_count += 1
            
            # If we still only have 1 page but the document is large, try a different approach
            if page_count == 1 and os.path.getsize(docx_path) > 102400:  # If file > 100KB
                # Estimate based on word count as fallback
                word_count = sum(len(para.text.split()) for para in doc.paragraphs)
                estimated_pages = max(1, word_count // 250)  # Rough estimate: 250 words per page
                print(f"Using word count fallback: {word_count} words ≈ {estimated_pages} pages")
                return estimated_pages
            
            print(f"Detected {page_count} pages in document")
            return max(1, page_count)  # Ensure at least 1 page
            
        except Exception as e:
            print(f"Warning: Could not count pages using page breaks: {e}")
            # Fallback to a simple word count estimate
            try:
                doc = Document(docx_path)
                word_count = sum(len(para.text.split()) for para in doc.paragraphs)
                estimated_pages = max(1, word_count // 250)  # Rough estimate: 250 words per page
                print(f"Using word count fallback: {word_count} words ≈ {estimated_pages} pages")
                return estimated_pages
            except:
                return default_pages
            
    except ImportError:
        print("Warning: python-docx not available, using default page count")
        return 1  # Default to 1 page if python-docx is not available
        
    except Exception as e:
        print(f"Warning: Could not count pages: {e}")
        return 1  # Fallback: Return a default of 1 page
        return 1

def filter_team_actions(action_items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Filter action items to only include those assigned to teams.
    Uses pandas for efficient filtering.
    
    Args:
        action_items: List of action item dictionaries
        
    Returns:
        Filtered list of action items assigned to teams
    """
    if not action_items:
        return []
    
    # Common team indicators (case insensitive)
    team_indicators = [
        'ABC', 'GBS'
    ]
    
    # Convert to DataFrame for easier manipulation
    df = pd.DataFrame(action_items)
    
    # Ensure 'responsible' column exists and handle missing values
    if 'responsible' not in df.columns:
        df['responsible'] = 'Unassigned'
    df['responsible'] = df['responsible'].fillna('Unassigned')
    
    # Filter for team assignments
    team_filter = df['responsible'].str.lower().str.contains(
        '|'.join(team_indicators), 
        case=False, 
        na=False
    )
    
    # Also include items where responsible is a known team name
    known_teams = [
        'ABC', 'GBS'
    ]
    team_name_filter = df['responsible'].str.lower().isin(known_teams)
    
    # Combine filters
    filtered_df = df[team_filter | team_name_filter]
    
    # Convert back to list of dictionaries
    return filtered_df.to_dict('records')

def highlight_chat_lines(text: str) -> str:
    """
    Highlight 'Shared the following in the chat' lines with yellow background
    and process only the first image.
    """
    # First process images (only first one)
    text_with_single_image = process_first_image_only_for_display(text)
    
    # Then highlight chat lines
    highlighted_text = re.sub(
        r'(Shared the following in the chat)',
        r'<mark style="background-color: yellow; color: black;">\1</mark>',
        text_with_single_image,
        flags=re.IGNORECASE
    )
    return highlighted_text

def process_single_image(match_obj, doc):
    """Process a single image match and add it to the document.
    
    Args:
        match_obj: The regex match object containing the image data
        doc: The Word document object to add the image to
        
    Returns:
        str: An empty string (since we're handling the image directly in the doc)
    """
    import base64
    import tempfile
    from docx.shared import Inches
    
    # Get the matched image text
    img_text = match_obj.group(0)
    
    # Process the image using the existing replace_png_placeholder_fixed function
    processed_img = replace_png_placeholder_fixed(match_obj)
    
    # Check if the processed result contains base64 image data
    base64_pattern = r'<img src="data:image/([^;]+);base64,([^"]+)"[^>]*>'
    base64_matches = re.findall(base64_pattern, processed_img)
    
    for img_format, base64_data in base64_matches:
        try:
            # Decode base64 data
            img_data = base64.b64decode(base64_data)
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{img_format}') as temp_file:
                temp_file.write(img_data)
                temp_img_path = temp_file.name
            
            # Add image to document
            doc.add_picture(temp_img_path, width=Inches(6.0))
            doc.add_paragraph()  # Add space after image
            
            # Clean up temporary file
            try:
                os.remove(temp_img_path)
            except:
                pass
                
        except Exception as e:
            print(f"Error processing base64 image: {e}")
            doc.add_paragraph("[Image could not be embedded]")
    
    # Return empty string as we've already added the image to the doc
    return ""

def create_summary_document(action_items, observations, summary_text="", output_path="final_summarized.docx"):
    """
    Create a Word document with action items and observations in a table format.
    Only includes actions assigned to teams. Processes HTML base64 images directly.
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Meeting Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    from datetime import datetime
    date_para = doc.add_paragraph()
    date_para.add_run(f"Summary Date: {datetime.now().strftime('%B %d, %Y')}").italic = True
    
    # Add summary text if provided
    if summary_text:
        doc.add_paragraph()
        # Add timeline header in bold
        timeline_para = doc.add_paragraph()
        timeline_para.add_run("Timelines (Times are in Eastern time (GMT-5) unless otherwise noted):").bold = True
        
        # Process summary text for base64 images
        import re
        import base64
        import tempfile
        import os
        from docx.shared import Inches
        
        def process_base64_images_in_text(text_content):
            """Process base64 images in HTML img tags and embed them at exact positions"""
            
            # Track if any image has been processed (only process the first image found)
            image_already_processed = False
            
            # Pattern to find HTML img tags with base64 data
            base64_pattern = r'<img src="data:image/([^;]+);base64,([^"]+)"[^>]*>'
            
            # Split text into parts around the first base64 image
            match = re.search(base64_pattern, text_content)
            
            if match and not image_already_processed:
                # Get text before and after the image
                before_image = text_content[:match.start()]
                after_image = text_content[match.end():]
                
                img_format = match.group(1)
                base64_data = match.group(2)
                
                try:
                    # Decode base64 data
                    img_data = base64.b64decode(base64_data)
                    
                    # Create temporary file
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{img_format}') as temp_file:
                        temp_file.write(img_data)
                        temp_img_path = temp_file.name
                    
                    # Add text before image
                    before_clean = re.sub(r'<[^>]+>', '', before_image).strip()
                    if before_clean:
                        for line in before_clean.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line)
                    
                    # Add the image at exact position
                    doc.add_picture(temp_img_path, width=Inches(6.0))
                    doc.add_paragraph()  # Add space after image
                    
                    # Clean up temporary file
                    try:
                        os.remove(temp_img_path)
                    except:
                        pass
                    
                    # Process remaining text (remove any subsequent images)
                    remaining_text = re.sub(base64_pattern, '', after_image)
                    remaining_clean = re.sub(r'<[^>]+>', '', remaining_text).strip()
                    if remaining_clean:
                        for line in remaining_clean.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line)
                    
                    image_already_processed = True
                    
                except Exception as e:
                    print(f"Error processing base64 image: {e}")
                    # Add text before the failed image
                    before_clean = re.sub(r'<[^>]+>', '', before_image).strip()
                    if before_clean:
                        for line in before_clean.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line)
                    
                    doc.add_paragraph("[Base64 image could not be embedded]")
                    
                    # Add text after the failed image
                    after_clean = re.sub(r'<[^>]+>', '', after_image).strip()
                    if after_clean:
                        for line in after_clean.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line)
            else:
                # No base64 images found, just clean HTML tags and add text
                clean_text = re.sub(r'<[^>]+>', '', text_content).strip()
                if clean_text:
                    for line in clean_text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)
        
        # Process the summary text with base64 image embedding at exact positions
        process_base64_images_in_text(summary_text)
 
    # Filter action items to only include team assignments
    team_action_items = filter_team_actions(action_items)
    
    # Add Action Items section
    doc.add_heading('Action Items', level=1)
    
    if team_action_items:
        # Create a table with 2 columns: Action and Team
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(5.5)  # Wider column for action
        table.columns[1].width = Inches(1.5)  # Narrower column for team
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Action Item'
        hdr_cells[1].text = 'Team'
        
        # Make headers bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for item in team_action_items:
            row_cells = table.add_row().cells
            action_text = item.get('action', '')
            team_name = item.get('responsible', 'Unassigned')
            
            # Add status as a note in the action cell if not 'Open'
            status = item.get('status', 'Open')
            if status.lower() != 'open':
                action_text += f"\n(Status: {status})"
            
            row_cells[0].text = action_text
            row_cells[1].text = team_name
    else:
        doc.add_paragraph("No team action items found.")
    
    # Add Observations section
    doc.add_paragraph()
    doc.add_heading('Key Observations', level=1)
    
    if observations:
        for obs in observations:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(obs).font.size = Pt(11)
    else:
        doc.add_paragraph("No key observations found.")
    
    # Save the document
    doc.save(output_path)
    return output_path
def main():
    """
    Launch Streamlit app for transcript summarization.
    """
    import re  # Add regex import for image processing
    
    # Set up the Streamlit page
    st.set_page_config(layout="wide", page_title="Meeting Summarizer")
    
    # Initialize COM for the main thread if on Windows
    if platform.system() == 'Windows' and not hasattr(threading.current_thread(), "_com_initialized"):
        try:
            pythoncom.CoInitialize()
            threading.current_thread()._com_initialized = True
        except Exception as e:
            st.warning(f"Warning: Could not initialize COM: {e}")
            if not hasattr(threading.current_thread(), "_com_initialized"):
                threading.current_thread()._com_initialized = False
    st.title("📝 Meeting Summarizer & Knowledge Base Generator")
    
    # Add application description
    st.markdown("""
    <style>
        @keyframes flash {
            0% { box-shadow: 0 0 0 0 rgba(0, 123, 255, 0.8); }
            50% { box-shadow: 0 0 0 10px rgba(0, 123, 255, 0); }
            100% { box-shadow: 0 0 0 0 rgba(0, 123, 255, 0); }
        }
        .flash-button {
            animation: flash 1s ease-out 5;
            position: relative;
            z-index: 1;
        }
        .description-container {
            background-color: #f8f9fa;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 20px;
        }
        .description-text {
            font-size: 16px;
            line-height: 1.6;
        }
        .steps-list {
            margin-left: 20px;
        }
        .steps-list li {
            margin-bottom: 8px;
        }
        .upload-container {
            border: 2px solid #e0e0e0;
            border-radius: 10px;
            padding: 20px;
            margin-bottom: 20px;
            background-color: #fafafa;
        }
        .merged-display {
            border: 2px solid #ccc;
            border-radius: 8px;
            padding: 15px;
            background-color: #f8f9fa;
            height: 300px;
            overflow-y: auto;
            font-family: monospace;
            font-size: 14px;
            white-space: pre-wrap;
        }
        .status-section {
            text-align: center;
            padding: 20px;
        }
        .generate-button {
            background-color: #0066cc;
            color: white;
            border: none;
            padding: 15px 30px;
            font-size: 18px;
            border-radius: 8px;
            cursor: pointer;
            margin: 20px 0;
        }
        .final-output {
            border: 2px solid #0066cc;
            border-radius: 8px;
            padding: 15px;
            background-color: #e6f3ff;
            height: 500px;
            overflow-y: auto;
            font-size: 16px;
            white-space: pre-line;
        }
        .qa-button-container {
            margin-top: 100px;
            text-align: center;
        }
    </style>
    
    <div class="description-container">
        <p class="description-text">
            Hi there! Ready to transform your team huddles into Smart & Actionable Summaries in a jiffy!
        </p>
        <p class="description-text">
            <strong>Steps to use the application:</strong>
        </p>
        <ul class="steps-list">
            <li>📁 Upload your meeting transcript and chat files</li>
            <li><span style="color: #D2B48C; text-shadow: 0 0 2px rgba(0,0,0,0.2);">👆</span> Click on Generate Whiteboard button</li>
            <li>📥 Download the summarized file</li>
            <li><span style="color: #D2B48C; text-shadow: 0 0 2px rgba(0,0,0,0.2);">👆</span> Under Q&A tab, click Extract Q&A button</li>
            <li>📥 Download the Q&A file</li>
        </ul>
    </div>
    """,
    unsafe_allow_html=True
    )
    
    # Create tabs
    tab1, tab2 = st.tabs(["📋 Summary", "❓ Q&A"])
    
    with tab1:
        # Initialize session state variables
        if 'page_count_msg' not in st.session_state:
            st.session_state.page_count_msg = 0

        if 'download_doc' not in st.session_state:
            st.session_state.download_doc = False
        
        st.header("Meeting Summarization")

        # Single file upload section for both transcript and chat files
        st.markdown("### File Upload")
        
        # Define file type detection function
        def get_file_type(filename):
            """Determine if the file is a transcript or chat based on filename."""
            filename_lower = filename.lower()
            if 'transcript' in filename_lower:
                return 'transcript'
            elif 'chat' in filename_lower:
                return 'chat'
            return None  # Unknown type
            
        uploaded_files = st.file_uploader(
            "📂 Upload Transcript and Chat Files",
            type=['docx'],
            accept_multiple_files=True,
            help="Upload your transcript and chat files (.docx). Files will be automatically detected as transcript or chat based on their filenames.",
            key="file_upload"
        )
        
        # Initialize variables
        transcript_file = None
        chat_file = None
        transcript_content = ""
        chat_content = ""
        
        # Process uploaded files
        if uploaded_files and len(uploaded_files) > 0:
            for file in uploaded_files:
                file_type = get_file_type(file.name)
                if file_type == 'transcript' and transcript_file is None:
                    transcript_file = file
                    
                    # Save the uploaded file temporarily to count pages
                    temp_path = os.path.join(tempfile.gettempdir(), file.name)
                    with open(temp_path, "wb") as f:
                        f.write(file.getbuffer())
                    
                    # Calculate and store total pages in session state
                    total_pages = count_pages_in_docx(temp_path)
                    st.session_state['total_pages'] = total_pages
                    print(f"Total pages in document: {total_pages}")  # Debug log
                    
                    # Clean up the temporary file
                    try:
                        os.remove(temp_path)
                    except:
                        pass
                elif file_type == 'chat' and chat_file is None:
                    chat_file = file
                
            # Show file status
            if transcript_file and chat_file:
                st.success("✅ Both transcript and chat files detected!")
            elif transcript_file:
                st.warning("ℹ️ Only transcript file detected. Please upload a chat file as well.")
            elif chat_file:
                st.warning("ℹ️ Only chat file detected. Please upload a transcript file as well.")
            else:
                st.error("❌ Could not determine file types. Please ensure filenames contain 'transcript' or 'chat'.")
        
        # Process files when both are uploaded
        if transcript_file is not None and chat_file is not None:
            try:
                # Save uploaded files temporarily
                transcript_temp_path = "temp_transcript.docx"
                chat_temp_path = "temp_chat.docx"
                merged_temp_path = "temp_merged.docx"
                
                with open(transcript_temp_path, "wb") as f:
                    f.write(transcript_file.getvalue())
                    
                with open(chat_temp_path, "wb") as f:
                    f.write(chat_file.getvalue())
                
                # Create status placeholder for file processing
                file_status_placeholder = st.empty()
                
                file_status_placeholder.info("📂 Loading transcript file...")
                # Load individual file contents for display
                transcript_content = load_text_from_docx(transcript_temp_path, encoding='utf-8-sig')
                
                file_status_placeholder.info("💬 Loading chat file...")
                chat_content = load_text_from_docx(chat_temp_path, encoding='utf-8-sig')
                
                file_status_placeholder.info("🔄 Merging transcript and chat files...")
                # Merge the files using merge_chat_transcript.py functions
                merge_chat_and_transcript(transcript_temp_path, chat_temp_path, merged_temp_path)
                
                file_status_placeholder.info("📋 Loading merged content...")
                # Load the merged text
                raw = load_text_from_docx(merged_temp_path, encoding='utf-8-sig')
                st.session_state["raw"] = raw
                st.session_state["transcript_content"] = transcript_content
                st.session_state["chat_content"] = chat_content
                st.session_state.pop("streamed", None)
                st.session_state.pop("final", None)
                st.session_state["temp_path"] = merged_temp_path
                
                file_status_placeholder.success("✅ Files merged and processed successfully!")
                
                # Clear status after 2 seconds (optional - you can remove this if you want it to stay)
                import time
                time.sleep(1)
                file_status_placeholder.empty()
                
                # Clean up individual temp files
                if os.path.exists(transcript_temp_path):
                    os.remove(transcript_temp_path)
                if os.path.exists(chat_temp_path):
                    os.remove(chat_temp_path)
                    
            except Exception as e:
                st.error(f"Error processing the files: {str(e)}")
                # Clean up temp files if they exist
                for temp_file in [transcript_temp_path, chat_temp_path, merged_temp_path]:
                    if 'temp_file' in locals() and os.path.exists(temp_file):
                        os.remove(temp_file)

        # Main layout with three columns (after file upload)
        left_col, middle_col, right_col = st.columns([2, 1, 2])

        with left_col:
            # Display transcript content
            st.markdown("### Transcript")
            if "transcript_content" in st.session_state:
                st.markdown(
                    f'<div class="merged-display">{st.session_state["transcript_content"]}</div>',
                    unsafe_allow_html=True
                )
            else:
                st.markdown(
                    '<div class="merged-display" style="display: flex; align-items: center; justify-content: center; color: #666;">'
                    'Transcript content will appear here'
                    '</div>', 
                    unsafe_allow_html=True
                )
            
            st.markdown("### Chat")
            if "chat_content" in st.session_state:
                # Process chat content to handle image placeholders
                chat_content = st.session_state["chat_content"]
                
                # Process image placeholders in the chat content
                if chat_content and isinstance(chat_content, str):
                    # Process all image placeholders
                    chat_content = re.sub(
                        r'\[(?:IMAGE|Image):\s*([^]]+?)(?:\s*\[Image:[^]]+\])?\]', 
                        replace_png_placeholder_fixed,
                        chat_content,
                        flags=re.IGNORECASE
                    )
                
                st.markdown(
                    f'<div class="merged-display">{chat_content}</div>',
                    unsafe_allow_html=True
                )
            else:
                st.markdown(
                    '<div class="merged-display" style="display: flex; align-items: center; justify-content: center; color: #666;">'
                    'Chat content will appear here'
                    '</div>', 
                    unsafe_allow_html=True
                )

        with middle_col:
            st.markdown('<div class="status-section">', unsafe_allow_html=True)
            
            # Status updates and button section
            if "raw" in st.session_state:
                # Generate Whiteboard button with margin to match tab2
                st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)
                if st.button("🚀 Generate Whiteboard", key="merge_button", use_container_width=True):
                    st.session_state.merge_button_clicked = True
                    
                    # Create persistent placeholders for status updates
                    status_placeholder = st.empty()
                    page_count_placeholder = st.empty()
                    progress_placeholder = st.empty()
                    page_status_placeholder = st.empty()
                    
                    # Initialize status
                    status_placeholder.info("🚀 Starting summary process...")
                    
                    raw = st.session_state["raw"]
                    temp_path = st.session_state.get("temp_path")
                    
                    # Load document and extract images
                    try:
                        # Create temp directory for images if it doesn't exist
                        temp_img_dir = os.path.join(os.path.dirname(temp_path), 'temp_images')
                        if not os.path.exists(temp_img_dir):
                            os.makedirs(temp_img_dir)
                            
                        # Load document
                        doc = Document(temp_path)
                        
                        # Create a new document to store the processed content
                        processed_doc = Document()
                        
                        # Process each paragraph in the document
                        for para in doc.paragraphs:
                            # Extract images from the current paragraph
                            image_paths = extract_images_from_paragraph(para, output_dir=temp_img_dir)
                            
                            # If there are images, add them as separate paragraphs
                            for img_path in image_paths:
                                img_para = processed_doc.add_paragraph()
                                img_para.add_run(f"[IMAGE:{img_path}]")
                            
                            # Add the paragraph text if it's not empty
                            if para.text.strip():
                                # Add the paragraph with the same style as the original
                                new_para = processed_doc.add_paragraph(para.text, style=para.style.name)
                                
                                # Copy paragraph formatting
                                new_para.alignment = para.alignment
                                new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                                new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                                new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                                new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
                                new_para.paragraph_format.space_before = para.paragraph_format.space_before
                                new_para.paragraph_format.space_after = para.paragraph_format.space_after
                        
                        # Save the processed document to a temporary file
                        processed_path = os.path.splitext(temp_path)[0] + "_processed.docx"
                        processed_doc.save(processed_path)
                        
                        # Update the document path to use the processed version
                        doc = Document(processed_path)
                        status_placeholder.info("📄 Document loaded and processed successfully...")
                        
                    except Exception as e:
                        status_placeholder.error(f"Error loading document: {str(e)}")
                        # Clean up temporary files if they exist
                        if 'processed_path' in locals() and os.path.exists(processed_path):
                            try:
                                os.remove(processed_path)
                            except:
                                pass
                        return
                    
                    # Extract meeting information
                    status_placeholder.info("🔍 Extracting meeting information...")
                    
                    paragraphs = list(iter_block_items(doc))
                    
                    if len(paragraphs) >= 2 and isinstance(paragraphs[1], Paragraph):
                        second_text = paragraphs[1].text.strip()
                        meeting_date, start_time = extract_meeting_datetime(second_text)
                        
                        if not meeting_date or not start_time:
                            # Try to extract from merged content header
                            lines = raw.split('\n')[:5]  # Check first 5 lines
                            for line in lines:
                                meeting_date, start_time = extract_meeting_datetime(line.strip())
                                if meeting_date and start_time:
                                    break
                            
                            if not meeting_date or not start_time:
                                status_placeholder.error("Could not extract meeting date and start time")
                                return
                    else:
                        status_placeholder.error("No valid meeting information found")
                        return
                    
                    # Process text content
                    status_placeholder.info("📝 Extracting text content...")
                    
                    all_text = ""
                    for block in iter_block_items(doc):
                        if isinstance(block, Paragraph):
                            text = block.text.strip()
                            text = text.replace('\n', '\\n')
                            all_text += text + "\n"
                    
                    # Split into chunks and process
                    status_placeholder.info("✂️ Preparing text for processing...")
                    chunks = _split_raw_into_chunks(all_text)
                    num_chunks = len(chunks)
                    
                    # Estimate total pages based on content length
                    total_words = len(all_text.split())
                    estimated_pages = max(1, total_words // 500)  # Roughly 500 words per page
                    pages_per_chunk = max(1, round(estimated_pages / len(chunks)))
                    
                    # Initialize progress bar
                    progress_bar = progress_placeholder.progress(0)
                    
                    # Initialize the final output container in the right column with header
                    with right_col:
                        st.markdown("### Final Summarized Output")
                        final_display = st.empty()
                        final_display.markdown(
                            "<div id='final_output' style='font-size:18px; height:800px; overflow-y:auto; padding:8px; border:1px solid #ccc; background-color: #e6f3ff; white-space: pre-line;'></div>",
                            unsafe_allow_html=True
                        )
                    
                    # Initialize the final output content
                    final_output_content = ""
                    
                    # Process chunks
                    processed_chunks = []
                    current_summary = ""
                    
                    # Initialize image processing flag for the entire summary
                    image_processed_in_summary = False
                    
                    for i, chunk in enumerate(chunks):
                        try:
                            # Get the summary from the generator
                            summary = next(chunked_clean_and_summarize(chunk))
                            
                            # Format timestamps
                            lines = summary.split('\n')
                            formatted_lines = []
                            for line in lines:
                                if not line.strip():
                                    continue
                                first_space = line.find(' ')
                                timestamp = line[:first_space]
                                full_date = format_timestamp(timestamp, meeting_date, start_time)
                                full_date = reformat_timestamp(full_date)
                                
                                if first_space == -1:
                                    formatted_line = f"{full_date} {line.title()}"
                                else:
                                    parts = line[first_space+1:].split(' ', 1)
                                    if len(parts) == 2:
                                        speaker = parts[0].title()
                                        content = parts[1]
                                        formatted_line = f"{full_date} {speaker} {content}"
                                    else:
                                        formatted_line = f"{full_date} {line[first_space+1:].title()}"
                                formatted_lines.append(formatted_line)
                            
                            formatted_chunk = '\n'.join(formatted_lines)
                            
                            # Process image placeholders - only process the first one across all chunks
                            if not image_processed_in_summary:
                                # Check if this chunk contains an image
                                if re.search(r'\[(?:IMAGE|Image):\s*([^]]+?)(?:\s*\[Image:[^]]+\])?\]', formatted_chunk, re.IGNORECASE):
                                    # Process only the first image
                                    formatted_chunk, num_subs = re.subn(
                                        r'\[(?:IMAGE|Image):\s*([^]]+?)(?:\s*\[Image:[^]]+\])?\]', 
                                        replace_png_placeholder_fixed,
                                        formatted_chunk,
                                        count=1,  # Only replace the first occurrence
                                        flags=re.IGNORECASE
                                    )
                                    if num_subs > 0:
                                        image_processed_in_summary = True
                                        print(f"\n=== DEBUG: Processed first image in chunk {i+1} ===")
                            
                            # Remove any remaining image placeholders from this chunk
                            formatted_chunk = re.sub(
                                r'\[(?:IMAGE|Image):[^\]]+\]', 
                                '',
                                formatted_chunk,
                                flags=re.IGNORECASE
                            )
                            
                            processed_chunks.append(formatted_chunk)
                            current_summary += formatted_chunk + "\n\n"
                            
                            # Update the final output content with the new section
                            final_output_content = current_summary
                            
                            # Update the final output container with the complete summary and header
                            final_display.empty()  # Clear previous content
                            final_display.markdown(
                                f"<div id='final_output' style='font-size:18px; height:800px; overflow-y:auto; padding:8px; border:1px solid #ccc; background-color: #e6f3ff; white-space: pre-line;'>{final_output_content}</div>",
                                unsafe_allow_html=True
                            )
                            
                            # Get total_pages from session state or default to 1
                            total_pages = st.session_state.get('total_pages', 1)
                            
                            # Calculate current page and progress
                            # Use min to ensure we don't exceed the total number of pages
                            current_page = min((i + 1) * pages_per_chunk, total_pages)
                            
                            # Ensure we don't show more than 100% progress
                            progress_percentage = min(1.0, (i + 1) / num_chunks)
                            progress_bar.progress(progress_percentage)
                            
                            # Show actual page count instead of estimated
                            status_placeholder.info(f"⚡ Processed {min((i + 1) * pages_per_chunk, total_pages)}/{total_pages} pages of transcript")
                            
                        except Exception as e:
                            status_placeholder.error(f"Error processing section {i}: {str(e)}")
                            break
                    
                    # Update the final status to show completion
                    page_status_placeholder.text(f"Processed {total_pages}/{total_pages} pages")
                    
                    # Final processing
                    status_placeholder.info("🔧 Finalizing summary and extracting insights...")
                    
                    try:
                        final_summary = "\n\n".join(processed_chunks)
                        
                        # Extract action items and observations
                        status_placeholder.info("📋 Extracting action items...")
                        action_items = extract_action_items(final_summary)
                        
                        status_placeholder.info("🔍 Identifying key observations...")
                        key_observations = extract_key_observations(final_summary)
                        priority_observations = filter_priority_observations(key_observations)
                        
                        # Create final output
                        team_action_items = filter_team_actions(action_items)
                        
                        final_output_with_actions = (
                            f"{final_summary}\n\n"
                            "Action Items:\n" + "\n".join(f"- {i['action']} ({i['responsible']})" for i in team_action_items) + "\n\n"
                            "Key Observations:\n" + "\n".join(f"- {j}" for j in priority_observations)
                        )
                        
                        # Apply single image processing to the final output
                        final_output_with_actions = process_first_image_only_for_display(final_output_with_actions)
                        
                        # Highlight chat lines
                        final_output_with_actions = re.sub(
                            r'(Shared the following in the chat)',
                            r'<mark style="background-color: yellow; color: black;">\1</mark>',
                            final_output_with_actions,
                            flags=re.IGNORECASE
                        )
                        
                        # Store in session state
                        st.session_state["streamed"] = final_summary
                        st.session_state["final"] = final_output_with_actions
                        st.session_state["final_summary_text"] = final_summary
                        st.session_state["action_items"] = action_items
                        st.session_state["key_observations"] = key_observations
                        st.session_state["priority_observations"] = priority_observations
                        st.session_state["final_output"] = final_output_with_actions
                        
                        # Update the final output container with the complete summary (with single image)
                        final_display.empty()  # Clear previous content
                        final_display.markdown(
                            f"<div id='final_output' style='font-size:18px; height:800px; overflow-y:auto; padding:8px; border:1px solid #ccc; background-color: #e6f3ff; white-space: pre-line;'>{final_output_with_actions}</div>"
                            "<script>"
                            "const element = document.getElementById('final_output');"
                            "element.scrollTop = element.scrollHeight;"
                            "</script>",
                            unsafe_allow_html=True
                        )
                        
                        # Create document
                        status_placeholder.info("💾 Creating downloadable document...")
                        doc_path = create_summary_document(
                            action_items=action_items,
                            observations=priority_observations,
                            summary_text=final_summary,
                            output_path="final_summarized.docx"
                        )
                        st.session_state["doc_path"] = doc_path
                        
                    except Exception as e:
                        status_placeholder.error(f"Error during final processing: {str(e)}")
                        return
                    
                    # Final success
                    progress_bar.progress(1.0)
                    status_placeholder.success("✅ Summary created successfully!")
                    
                    # Clean up temp file
                    if temp_path and os.path.exists(temp_path):
                        os.remove(temp_path)
                        st.session_state.pop("temp_path", None)
                    
                    st.rerun()
                
                # Download button (shown after processing)
                if "doc_path" in st.session_state and os.path.exists(st.session_state["doc_path"]):
                    with open(st.session_state["doc_path"], "rb") as file:
                        st.download_button(
                            label="📥 Download Whiteboard",
                            data=file,
                            file_name="final_summarized.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="download_final"
                        )
                
                # If we're just displaying from session state (after page refresh)
                if "final_output" in st.session_state and "merge_button_clicked" not in st.session_state:
                    # Get total_pages from session state or default to 1
                    total_pages = st.session_state.get("total_pages", 1)
                    # Initialize placeholders
                    status_placeholder = st.empty()
                    progress_placeholder = st.empty()
                    page_count_placeholder = st.empty()
                    
                    progress_bar = progress_placeholder.progress(1.0)
                    status_placeholder.success("✅ Summary created successfully!")
                    
                    if "doc_path" in st.session_state and os.path.exists(st.session_state["doc_path"]):
                        with open(st.session_state["doc_path"], "rb") as file:
                            st.download_button(
                                label="📥 Download Whiteboard",
                                data=file,
                                file_name="final_summarized.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="download_final_refresh"
                            )
            else:
                st.info("Upload files to get started")
            
            st.markdown('</div>', unsafe_allow_html=True)

        with right_col:
            st.markdown("### Final Summarized Output")
            
            # Show merged content first (before processing)
            if "raw" in st.session_state and "final_output" not in st.session_state:
                # Process the raw content to show only the first image
                processed_raw = process_first_image_only_for_display(st.session_state["raw"])
                st.markdown(
                    f'<div class="final-output" style="height: 300px;">{processed_raw}</div>',
                    unsafe_allow_html=True
                )
                st.markdown("---")
            
            # Show processed summary (after processing)
            if "final_output" in st.session_state:
                # Ensure single image processing is applied
                processed_final_output = process_first_image_only_for_display(st.session_state["final_output"])
                st.markdown(
                    f'<div id="final_output" style="font-size:18px; height:670px; overflow-y:auto; padding:16px; border:1px solid #ccc; border-radius: 10px; background-color: #e6f3ff; white-space: pre-line; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">{processed_final_output}</div>',
                    unsafe_allow_html=True
                )
            elif "raw" not in st.session_state:
                # Show placeholder when no files uploaded
                st.markdown(
                    '<div class="final-output" style="display: flex; align-items: center; justify-content: center; color: #666;">'
                    'Upload files and generate summary to see results here'
                    '</div>',
                    unsafe_allow_html=True
                )

    # Q&A Tab (keeping the existing Q&A tab code)
    with tab2:
        st.header("Question & Answer Extractor")
        
        # Check if summary is available
        if "final_summary_text" not in st.session_state:
            st.info("📋 Please generate a summary first in the Summary tab to extract Q&A.")
        else:
            # Create columns for Q&A tab
            qa_left_col, qa_middle_col, qa_right_col = st.columns([1, 0.5, 1])
            
            with qa_left_col:
                st.subheader("Summary Text")
                # Apply single image processing to summary text display
                summary_text = process_first_image_only_for_display(st.session_state["final_summary_text"])
                st.markdown(
                    f"<div style='font-size:16px; white-space: pre-wrap; "
                    f"height:600px; overflow:auto; border:1px solid #ccc; padding:8px; background-color: #f8f9fa;border-radius: 10px;box-shadow: 0 2px 4px rgba(0,0,0,0.1);'>"
                    f"{summary_text}</div>",
                    unsafe_allow_html=True,
                )
            
            with qa_middle_col:
                # Add margin-top to align button with container start
                st.markdown('<div class="qa-button-container">', unsafe_allow_html=True)
                
                if st.button("🔍 Extract Q&A", key="extract_qa_button", use_container_width=True):
                    # Create placeholders for status updates
                    qa_status_placeholder = st.empty()
                    qa_progress_placeholder = st.empty()
                    
                    qa_status_placeholder.info("🔍 Extracting questions and answers...")
                    qa_progress_bar = qa_progress_placeholder.progress(0)
                    
                    try:
                        # Extract FAQs from the summary
                        faqs = extract_faqs(st.session_state["final_summary_text"])
                        
                        qa_progress_bar.progress(0.8)
                        qa_status_placeholder.info("📝 Creating FAQ document...")
                        
                        if faqs:
                            # Create FAQ document
                            faq_doc_path = create_faq_document(faqs, "extracted_faqs.docx")
                            st.session_state["faqs"] = faqs
                            st.session_state["faq_doc_path"] = faq_doc_path
                            qa_status_placeholder.info("📝 Created FAQ document.")
                            qa_progress_bar.progress(1.0)
                        else:
                            qa_progress_bar.progress(1.0)
                            qa_status_placeholder.warning("⚠️ No Q&A pairs found in the summary.")
                            st.session_state["faqs"] = []
                            
                    except Exception as e:
                        qa_status_placeholder.error(f"❌ Error extracting Q&A: {str(e)}")
                        print(f"FAQ extraction error: {e}")
                        import traceback
                        traceback.print_exc()
                
                # Show number of Q&A pairs extracted
                if "faqs" in st.session_state and st.session_state["faqs"]:
                    st.info(f"ℹ️ Extracted {len(st.session_state['faqs'])} Q&A pairs from the summary.")
                
                # Download button for FAQ document
                if "faq_doc_path" in st.session_state and os.path.exists(st.session_state["faq_doc_path"]):
                    with open(st.session_state["faq_doc_path"], "rb") as file:
                        st.download_button(
                            label="📥 Download Q&A Document",
                            data=file,
                            file_name="extracted_faqs.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_faq",
                            use_container_width=True
                        )
                
                st.markdown('</div>', unsafe_allow_html=True)
            
            with qa_right_col:
                st.subheader("Extracted Q&A")
                
                if "faqs" in st.session_state:
                    faqs = st.session_state["faqs"]
                    
                    if faqs:
                        # Display Q&A pairs
                        qa_display_content = ""
                        for i, faq in enumerate(faqs, 1):
                            question = faq.get('question', 'No question provided')
                            answer = faq.get('answer', 'No answer provided')
                            qa_display_content += f"Q{i}: {question}\n\nA{i}: {answer}\n\n" + "="*50 + "\n\n"
                        
                        st.markdown(
                            f"<div style='font-size:16px; white-space: pre-wrap; "
                            f"height:600px; overflow:auto; border:1px solid #ccc; padding:8px; background-color: #e8f5e8;border-radius: 10px;box-shadow: 0 2px 4px rgba(0,0,0,0.1);'>"
                            f"{qa_display_content}</div>",
                            unsafe_allow_html=True,
                        )
                    else:
                        st.markdown(
                            "<div style='font-size:16px; height:600px; border:1px solid #ccc; padding:8px; background-color: #fff3cd; display: flex; align-items: center; justify-content: center;border-radius: 10px;box-shadow: 0 2px 4px rgba(0,0,0,0.1);'>"
                            "<p style='text-align: center; color: #856404;'>No Q&A pairs extracted from the summary.</p>"
                            "</div>",
                            unsafe_allow_html=True,
                        )
                else:
                    st.markdown(
                        "<div style='font-size:16px; height:600px; border:1px solid #ccc; padding:8px; background-color: #f8f9fa; display: flex; align-items: center; justify-content: center;border-radius: 10px;box-shadow: 0 2px 4px rgba(0,0,0,0.1);'>"
                        "<p style='text-align: center; color: #6c757d;'>Click 'Extract Q&A' to generate questions and answers.</p>"
                        "</div>",
                        unsafe_allow_html=True,
                    )

if __name__ == "__main__":
    try:
        main()
    finally:
        # Clean up COM if it was initialized in this thread
        if hasattr(threading.current_thread(), "_com_initialized"):
            pythoncom.CoUninitialize()
            delattr(threading.current_thread(), "_com_initialized")
