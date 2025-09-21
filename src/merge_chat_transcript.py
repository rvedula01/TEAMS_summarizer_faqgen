import os
from openai import OpenAI
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from datetime import datetime, timedelta
import re
from typing import List, Tuple, Dict, Any, Optional
import io
from PIL import Image

# Configuration
DEBUG = True

def debug_print(*args, **kwargs):
    """Print debug messages if DEBUG is True."""
    if DEBUG:
        print(*args, **kwargs)

# ---------- File I/O Helpers ----------

def read_docx(filepath: str) -> str:
    """Read text content from a DOCX file."""
    try:
        doc = Document(filepath)
        return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        raise Exception(f"Error reading {filepath}: {str(e)}")

def write_docx(text: str, filepath: str, images_data: List[Tuple[str, bytes]] = None) -> None:
    """Write text to a DOCX file with embedded images, splitting paragraphs on double newlines."""
    from docx import Document
    from docx.shared import Inches
    
    doc = Document()
    paragraphs = text.split('\n\n')
    image_index = 0
    
    for para in paragraphs:
        # Check if this paragraph contains an image placeholder
        if "[IMAGE_PLACEHOLDER_" in para and images_data and image_index < len(images_data):
            # Add the text part before the image
            text_part = re.sub(r'\[IMAGE_PLACEHOLDER_\d+\]', '', para).strip()
            if text_part:
                doc.add_paragraph(text_part)
            
            # Add the image
            try:
                image_name, image_data = images_data[image_index]
                # Create a temporary file-like object from bytes
                image_stream = io.BytesIO(image_data)
                
                # Add image to document with reasonable size
                paragraph = doc.add_paragraph()
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                
                # Try to determine appropriate size
                try:
                    with Image.open(io.BytesIO(image_data)) as img:
                        width, height = img.size
                        # Scale to fit reasonable document width (max 6 inches)
                        max_width = 6.0
                        if width > height:
                            new_width = min(max_width, width / 96)  # Convert pixels to inches (96 DPI)
                            new_height = (height / width) * new_width
                        else:
                            new_height = min(max_width, height / 96)
                            new_width = (width / height) * new_height
                        
                        run.add_picture(image_stream, width=Inches(new_width), height=Inches(new_height))
                except Exception as e:
                    debug_print(f"⚠️ Could not resize image {image_name}: {e}. Using default size.")
                    run.add_picture(image_stream, width=Inches(4))
                
                image_index += 1
                debug_print(f"✅ Embedded image: {image_name}")
                
            except Exception as e:
                debug_print(f"⚠️ Failed to embed image {image_index}: {e}")
                doc.add_paragraph(f"[Failed to embed image: {e}]")
                image_index += 1
        else:
            # Regular text paragraph
            doc.add_paragraph(para)
    
    try:
        doc.save(filepath)
    except Exception as e:
        raise Exception(f"Error writing to {filepath}: {str(e)}")

def write_text_file(text: str, filepath: str) -> None:
    """Write text to a plain text file."""
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(text)
    except Exception as e:
        raise Exception(f"Error writing to {filepath}: {str(e)}")

def extract_images_from_docx(doc_path: str) -> List[Tuple[str, bytes]]:
    """Extract all images from a DOCX file and return as list of (name, bytes) tuples."""
    images_data = []
    
    try:
        doc = Document(doc_path)
        
        # Access the document part to get relationships
        document_part = doc.part
        
        for rel in document_part.rels.values():
            if "image" in rel.target_ref:
                try:
                    # Get image data
                    image_part = rel.target_part
                    image_data = image_part.blob
                    image_name = rel.target_ref.split('/')[-1]  # Get filename from path
                    
                    images_data.append((image_name, image_data))
                    debug_print(f"✅ Extracted image: {image_name} ({len(image_data)} bytes)")
                    
                except Exception as e:
                    debug_print(f"⚠️ Failed to extract image from {rel.target_ref}: {e}")
                    
    except Exception as e:
        debug_print(f"⚠️ Error extracting images: {e}")
    
    debug_print(f"Total images extracted: {len(images_data)}")
    return images_data

# ---------- Enhanced Chat Format Conversion with Image Embedding ----------

def parse_chat_with_images(doc_path: str) -> Tuple[str, List[Tuple[str, bytes]]]:
    """
    Parse chat with proper image detection and extraction using document XML structure.
    Returns formatted text with image placeholders and extracted image data.
    """
    doc = Document(doc_path)
    images_data = extract_images_from_docx(doc_path)

    # Speaker line regex
    speaker_re = re.compile(
        r"(?P<name>.+?)(?:\s*\(External\))?\s*(?P<date>\d+/\d+)\s+(?P<time>\d{1,2}:\d{2})(?:Edited)?"
    )

    # Output
    chat_output = []
    current_speaker = None
    current_time = None
    message_buffer = []
    image_counter = 0

    def flush(force=False):
        nonlocal current_speaker, current_time
        if current_speaker and (message_buffer or force):
            content = ' '.join(message_buffer).strip() if message_buffer else ''
            chat_output.append(f"[Chat {current_time}:00] {current_speaker}: {content}")
            message_buffer.clear()
            if force and not content:
                # Reset speaker if nothing was actually said
                current_speaker = None
                current_time = None

    # Iterate over doc elements in true order
    for block in doc.element.body:
        if isinstance(block, CT_P):
            para = Paragraph(block, doc)

            # Check for image
            if 'graphic' in block.xml:
                if message_buffer and current_speaker and current_time:
                    # If there is a pending message, flush it and append image placeholder
                    content = ' '.join(message_buffer).strip()
                    chat_output.append(f"[Chat {current_time}:00] {current_speaker}: {content} [IMAGE_PLACEHOLDER_{image_counter}]")
                    message_buffer.clear()
                elif current_speaker and current_time:
                    # If no message, but speaker/time exist, output speaker line with image
                    chat_output.append(f"[Chat {current_time}:00] {current_speaker}: [IMAGE_PLACEHOLDER_{image_counter}]")
                elif chat_output and chat_output[-1].strip():
                    # Otherwise, try to append to previous line
                    chat_output[-1] = chat_output[-1].rstrip() + f" [IMAGE_PLACEHOLDER_{image_counter}]"
                else:
                    # If nothing to append to, treat as its own line
                    chat_output.append(f"[IMAGE_PLACEHOLDER_{image_counter}]")
                
                image_counter += 1
                # After image, reset speaker and time so next speaker starts new line
                current_speaker = None
                current_time = None
                message_buffer.clear()
                continue

            # Otherwise, process as paragraph
            text = para.text.strip()
            if not text:
                continue

            matches = list(speaker_re.finditer(text))
            if not matches:
                message_buffer.append(text)
                continue

            last_end = 0
            for match in matches:
                flush(force=True)
                pre_text = text[last_end:match.start()].strip()
                if pre_text:
                    message_buffer.append(pre_text)
                    flush()

                current_speaker = match.group("name").strip()
                current_time = match.group("time").strip()
                last_end = match.end()

            remaining = text[last_end:].strip()
            if remaining:
                message_buffer.append(remaining)

    flush(force=True)
    formatted_text = "\n\n".join(chat_output)
    
    # Only return the images that were actually referenced
    used_images = images_data[:image_counter] if image_counter <= len(images_data) else images_data
    
    return formatted_text, used_images

def extract_chat_entries_with_timestamps(formatted_chat: str, base_time: Optional[datetime] = None) -> List[Tuple[datetime, str]]:
    """
    Extract chat entries with proper datetime objects for chronological merging.
    """
    chat_entries = []
    
    if base_time is None:
        base_time = datetime.now().replace(second=0, microsecond=0)
        debug_print(f"⚠️ Using current time as base for chat: {base_time}")
    
    chat_lines = [line for line in formatted_chat.split('\n\n') if line.strip()]
    
    for line in chat_lines:
        # Match chat timestamp pattern [Chat HH:MM:SS]
        match = re.search(r"\[Chat (\d{1,2}):(\d{2}):(\d{2})\]", line)
        if match:
            h, m, s = map(int, match.groups())
            # Use base_time date with extracted time
            chat_time = base_time.replace(hour=h, minute=m, second=s)
            chat_entries.append((chat_time, line))
            debug_print(f"[{chat_time.strftime('%Y-%m-%d %H:%M:%S')}] Chat entry parsed")
        else:
            # Fallback for entries without proper timestamps
            chat_entries.append((base_time, line))
    
    debug_print(f"Extracted {len(chat_entries)} chat entries")
    return chat_entries

# ---------- Transcript Processing ----------

def extract_transcript_base_time(text: str) -> Optional[datetime]:
    """Extract the base timestamp from transcript header."""
    lines = text.split('\n')
    
    # Look for patterns like "June 24, 2025, 3:20PM"
    time_patterns = [
        r'([A-Za-z]+ \d{1,2}, \d{4}, \d{1,2}:\d{2}[AP]M)',
        r'(\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2})',
        r'(\d{1,2}/\d{1,2}/\d{4} \d{1,2}:\d{2}[AP]M)'
    ]
    
    formats = [
        "%B %d, %Y, %I:%M%p",
        "%Y-%m-%d %H:%M:%S", 
        "%m/%d/%Y %I:%M%p"
    ]
    
    for line in lines[:10]:  # Check first 10 lines
        for i, pattern in enumerate(time_patterns):
            match = re.search(pattern, line.strip())
            if match:
                try:
                    base_time = datetime.strptime(match.group(1), formats[i])
                    debug_print(f"✅ Found transcript base time: {base_time}")
                    return base_time
                except ValueError as e:
                    debug_print(f"⚠️ Failed to parse timestamp '{match.group(1)}': {e}")
                    continue
    
    debug_print("⚠️ No transcript header timestamp found")
    return None

def parse_time_offset(offset_str: str) -> int:
    """Parse time offset in MM:SS or HH:MM:SS format and return total seconds."""
    try:
        time_parts = offset_str.split(':')
        
        if len(time_parts) == 2:  # MM:SS format
            minutes, seconds = map(int, time_parts)
            return minutes * 60 + seconds
        elif len(time_parts) == 3:  # HH:MM:SS format
            hours, minutes, seconds = map(int, time_parts)
            return hours * 3600 + minutes * 60 + seconds
        else:
            raise ValueError(f"Invalid time format: {offset_str}")
            
    except ValueError as e:
        debug_print(f"⚠️ Failed to parse time offset '{offset_str}': {e}")
        return 0

def extract_transcript_segments(text: str) -> List[Tuple[datetime, str]]:
    """Extract transcript segments with timestamps calculated from base time + offset."""
    segments = []
    base_time = extract_transcript_base_time(text)
    
    if base_time is None:
        base_time = datetime.now().replace(second=0, microsecond=0)
        debug_print(f"⚠️ Using current time as base: {base_time}")

    lines = text.split('\n')
    current_speaker = None
    current_segment = []
    current_time = None
    current_offset_str = None  # store the original offset string for formatting

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Match speaker line like "Aryan   0:29" or "Aryan   1:02:40" (offset in MM:SS or HH:MM:SS)
        speaker_match = re.match(r'^([\w\s/()]+?)\s+(\d{1,2}:\d{2}(?::\d{2})?)\s*$', line)
        if speaker_match:
            # Save previous segment
            if current_segment and current_speaker and current_time and current_offset_str:
                message_text = '\n'.join(current_segment)
                segments.append((current_time, f"{current_speaker}   {current_offset_str}\n {message_text}"))
                current_segment = []

            current_speaker = speaker_match.group(1).strip()
            offset_str = speaker_match.group(2).strip()
            current_offset_str = offset_str

            # Parse time offset (handles both MM:SS and HH:MM:SS)
            total_seconds = parse_time_offset(offset_str)
            current_time = base_time + timedelta(seconds=total_seconds)
            debug_print(f"[{current_time.strftime('%Y-%m-%d %H:%M:%S')}] {current_speaker} (offset: {offset_str})")

        elif current_speaker:
            current_segment.append(line)

    # Add final segment
    if current_segment and current_speaker and current_time and current_offset_str:
        message_text = '\n'.join(current_segment)
        segments.append((current_time, f"{current_speaker}   {current_offset_str}\n {message_text}"))

    debug_print(f"Extracted {len(segments)} transcript segments")
    return segments
def reformat_merged_output(text: str, base_time: Optional[datetime] = None) -> str:
    """
    Reformat [Chat HH:MM:SS] entries into:
    Speaker   Time
    Shared the following in the chat:
    Message
    """
    lines = text.strip().split('\n')
    output_lines = []

    chat_pattern = re.compile(r'\[Chat (\d{1,2}):(\d{2}):(\d{2})\] (.+?):\s?(.*)')

    if base_time is None:
        base_time = datetime.now().replace(second=0, microsecond=0)

    for line in lines:
        match = chat_pattern.match(line.strip())
        if match:
            hour, minute, second, speaker, content = match.groups()
            chat_time = base_time.replace(hour=int(hour), minute=int(minute), second=int(second))
            offset_seconds = int((chat_time - base_time).total_seconds())
            hours, remainder = divmod(offset_seconds, 3600)
            minutes, seconds = divmod(remainder, 60)

            if hours == 0:
                if seconds == 0:
                    offset_str = f"{minutes}:00"
                else:
                    offset_str = f"{minutes}:{seconds:02}"
            else:
                offset_str = f"{hours}:{minutes:02}:{seconds:02}"

            formatted_block = f"{speaker}   {offset_str}\nShared the following in the chat:\n{content}"
            output_lines.append(formatted_block.strip())
        else:
            output_lines.append(line.strip())

    return '\n'.join(output_lines)

# ---------- Merging Functions ----------

def merge_chronologically(transcript_segments: List[Tuple[datetime, str]], 
                         chat_entries: List[Tuple[datetime, str]]) -> str:
    """Merge chat entries and transcript segments in chronological order."""
    
    # Combine all entries
    all_entries = transcript_segments + chat_entries
    
    # Sort by timestamp
    all_entries.sort(key=lambda x: x[0])
    
    debug_print(f"\nMerged timeline with {len(all_entries)} total entries:")
    for i, (ts, content) in enumerate(all_entries[:5]):  # Show first 5 entries
        debug_print(f"  {i+1}. [{ts.strftime('%H:%M:%S')}] {content[:100]}...")
    
    # Return merged content
    return '\n\n'.join([entry[1] for entry in all_entries])

# ---------- Main Functions ----------

def convert_chat_only(input_path: str, output_path: str) -> None:
    """Convert chat format only without merging, using enhanced image detection."""
    try:
        debug_print(f"Reading and parsing chat file: {input_path}")
        
        # Use enhanced chat parsing with image extraction
        formatted_chat, images_data = parse_chat_with_images(input_path)
        
        # Determine output format based on file extension
        if output_path.endswith('.docx'):
            write_docx(formatted_chat, output_path, images_data)
        else:
            # For text files, replace image placeholders with descriptive text
            text_output = re.sub(r'\[IMAGE_PLACEHOLDER_\d+\]', '[Image]', formatted_chat)
            write_text_file(text_output, output_path)
        
        print(f"✅ Chat conversion completed: {output_path}")
        if images_data:
            print(f"   📷 {len(images_data)} images embedded")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        raise

def merge_chat_and_transcript(transcript_path: str, chat_path: Optional[str] = None, output_path: str = None) -> str:
    """
    Merge chat and transcript files chronologically using enhanced chat parsing.
    If chat_path is None, only the transcript will be processed.
    
    Args:
        transcript_path: Path to the transcript file (required)
        chat_path: Path to the chat file (optional)
        output_path: Path to save the merged output. If None, a temporary file will be created.
        
    Returns:
        Path to the merged output file
    """
    try:
        debug_print("\n" + "="*50)
        debug_print("STARTING MERGE PROCESS")
        debug_print("="*50)

        # 1. Read and process transcript
        debug_print("\nReading transcript file...")
        with open(transcript_path, 'rb') as f:
            transcript_text = f.read().decode('utf-8', errors='ignore')
        
        debug_print("Extracting base time from transcript...")
        base_time = extract_transcript_base_time(transcript_text)
        
        debug_print("Extracting transcript segments...")
        transcript_segments = extract_transcript_segments(transcript_text)
        debug_print(f"Found {len(transcript_segments)} transcript segments")

        # If no chat file provided, just process the transcript
        if not chat_path:
            debug_print("\nNo chat file provided, processing transcript only...")
            # Create output path if not provided
            if not output_path:
                import tempfile
                _, output_path = tempfile.mkstemp(suffix='.docx')
            
            # Save the transcript directly
            with open(transcript_path, 'rb') as src, open(output_path, 'wb') as dst:
                dst.write(src.read())
            
            debug_print("\n" + "="*50)
            debug_print("TRANSCRIPT PROCESSING COMPLETED")
            debug_print("="*50)
            
            return output_path

        # 2. Read and process chat
        debug_print("\nProcessing chat file...")
        chat_text, chat_images = parse_chat_with_images(chat_path)
        chat_entries = extract_chat_entries_with_timestamps(chat_text, base_time)
        debug_print(f"Found {len(chat_entries)} chat entries")

        # 3. Merge chronologically
        debug_print("\nMerging content chronologically...")
        merged_entries = merge_chronologically(transcript_segments, chat_entries)
        
        # 4. Reformat and save
        debug_print("Reformatting output...")
        merged_text = '\n\n'.join(entry[1] for entry in merged_entries)
        reformatted_text = reformat_merged_output(merged_text, base_time)
        
        # 5. Save the merged document with images
        debug_print("Saving merged document...")
        if not output_path:
            import tempfile
            _, output_path = tempfile.mkstemp(suffix='.docx')
        
        write_docx(reformatted_text, output_path, chat_images)
        
        debug_print("\n" + "="*50)
        debug_print("MERGE COMPLETED SUCCESSFULLY")
        debug_print("="*50)
        
        return output_path
        
    except Exception as e:
        debug_print("\n" + "!"*50)
        debug_print(f"ERROR: {str(e)}")
        debug_print("!"*50)
        raise

def main():
    """Main function with enhanced command-line interface."""
    import sys
    
    if len(sys.argv) < 3:
        print("Usage:")
        print("  Convert chat only:")
        print("    python merge_chat_transcript.py convert <chat.docx> <output.txt|output.docx>")
        print("  Merge chat and transcript (chat is optional):")
        print("    python merge_chat_transcript.py merge <transcript.docx> [chat.docx] <output.docx>")
        print("\nNote: For merging, output must be a .docx file to support images.")
        sys.exit(1)
    
    command = sys.argv[1].lower()
    
    try:
        if command == 'convert':
            if len(sys.argv) != 4:
                raise ValueError("Incorrect number of arguments for convert command")
            
            input_path = sys.argv[2]
            output_path = sys.argv[3]
            
            print(f"Converting chat file: {input_path}")
            print(f"Output will be saved to: {output_path}")
            
            convert_chat_only(input_path, output_path)
            
        elif command == 'merge':
            # Check if the last argument is the output file (must be .docx)
            if not sys.argv[-1].lower().endswith('.docx'):
                raise ValueError("Output file must be a .docx file to support images")
                
            output_path = sys.argv[-1]
            transcript_path = sys.argv[2]
            
            # Check if a chat file was provided
            if len(sys.argv) == 5:  # merge transcript chat output
                chat_path = sys.argv[3]
                print(f"Merging transcript: {transcript_path}")
                print(f"With chat: {chat_path}")
                print(f"Output will be saved to: {output_path}")
                
                merge_chat_and_transcript(transcript_path, chat_path, output_path)
            else:  # merge transcript output only
                print(f"Processing transcript only: {transcript_path}")
                print(f"Output will be saved to: {output_path}")
                
                merge_chat_and_transcript(transcript_path, None, output_path)
            
        else:
            raise ValueError(f"Unknown command: {command}")
            
        print("✅ Operation completed successfully!")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        if DEBUG:
            import traceback
            traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()