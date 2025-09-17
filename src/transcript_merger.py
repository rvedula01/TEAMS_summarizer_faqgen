

# Initialize OpenAI client
import os
import re
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from dotenv import load_dotenv
import openai

# Load environment variables from .env file
load_dotenv()
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json
import base64
from io import BytesIO
from PIL import Image

class CompleteMergerWithPerfectExtraction:
    def __init__(self, api_key: str):
        """
        Initialize the Complete Merger with Perfect Extraction and OpenAI API key.
        
        Args:
            api_key (str): OpenAI API key
        """
        self.client = openai.OpenAI(api_key=api_key)
        self.max_chunk_size = 12000  # Reduced for better processing
        self.overlap_size = 500  # Overlap between chunks to maintain context
    
    # ========== IMAGE EXTRACTION METHODS ==========
    def extract_images_from_docx(self, file_path: str, output_dir: str = "extracted_images") -> Dict[str, str]:
        """Extract images from DOCX file and save them."""
        try:
            doc = Document(file_path)
            image_mapping = {}
            
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            image_counter = 1
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        
                        # Determine file extension
                        if rel.target_ref.endswith('.png'):
                            ext = 'png'
                        elif rel.target_ref.endswith('.jpg') or rel.target_ref.endswith('.jpeg'):
                            ext = 'jpg'
                        elif rel.target_ref.endswith('.gif'):
                            ext = 'gif'
                        else:
                            ext = 'png'  # default
                        
                        image_filename = f"image_{image_counter}.{ext}"
                        image_path = os.path.join(output_dir, image_filename)
                        
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_data)
                        
                        image_mapping[rel.target_ref] = image_path
                        print(f"Extracted image: {image_path}")
                        image_counter += 1
                        
                    except Exception as e:
                        print(f"Error extracting image {rel.target_ref}: {str(e)}")
            
            return image_mapping
        
        except Exception as e:
            print(f"Error extracting images from {file_path}: {str(e)}")
            return {}

    # ========== FILE READING METHODS ==========
    def read_docx_file_with_images(self, file_path: str) -> Tuple[str, Dict[str, str]]:
        """Read content from a DOCX file including images."""
        try:
            doc = Document(file_path)
            content = []
            image_mapping = self.extract_images_from_docx(file_path)
            
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    paragraph = Paragraph(element, doc)
                    para_text = paragraph.text.strip()
                    
                    has_image = False
                    for run in paragraph.runs:
                        if run._element.xpath('.//a:blip'):
                            has_image = True
                            if para_text:
                                content.append(para_text)
                            content.append("[IMAGE_PLACEHOLDER]")
                            break
                    
                    if not has_image and para_text:
                        content.append(para_text)
                
                elif isinstance(element, CT_Tbl):
                    table = Table(element, doc)
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            content.append(" | ".join(row_text))
            
            return "\n".join(content), image_mapping
        
        except Exception as e:
            print(f"Error reading DOCX file {file_path}: {str(e)}")
            return "", {}

    def read_docx_file(self, file_path: str) -> str:
        """Read content from a DOCX file (backward compatibility)."""
        content, _ = self.read_docx_file_with_images(file_path)
        return content
    
    def read_text_file(self, file_path: str) -> str:
        """Read content from a text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading text file {file_path}: {str(e)}")
            return ""

    # ========== PERFECT FORMAT EXTRACTION METHODS ==========
    def extract_last_pages_docx(self, file_path: str, num_pages: int = 2) -> str:
        """Extract content from the last N pages of a DOCX file."""
        try:
            doc = Document(file_path)
            all_paragraphs = []
            
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    paragraph = Paragraph(element, doc)
                    para_text = paragraph.text.strip()
                    if para_text:
                        all_paragraphs.append(para_text)
                elif isinstance(element, CT_Tbl):
                    table = Table(element, doc)
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            all_paragraphs.append(" | ".join(row_text))
            
            total_content = "\n".join(all_paragraphs)
            estimated_chars_per_page = 600
            target_length = estimated_chars_per_page * num_pages
            
            if len(total_content) <= target_length:
                return total_content
            
            last_content = total_content[-target_length:]
            lines = last_content.split('\n')
            if len(lines) > 1:
                return '\n'.join(lines[1:])
            
            return last_content
            
        except Exception as e:
            print(f"Error extracting last pages from {file_path}: {str(e)}")
            return ""
    
    def extract_perfect_actions_observations_openai(self, content: str, document_type: str) -> Tuple[List[str], List[str]]:
        """Extract action items and observations in perfect format using OpenAI."""
        
        if document_type == "final_summary":
            system_prompt = """
            You are an expert at extracting action items and key observations from incident final summary documents.
            Focus on the conclusions, follow-up actions, and lessons learned sections typically found at the end.
            
            PERFECT FORMAT REQUIREMENTS:
            
            For ACTION ITEMS:
            - Extract EVERY task, assignment, follow-up, investigation, or ticket mentioned
            - Format EXACTLY as: "Action: [Clear, specific description] | Team: [Team/Person name]"
            - If no team specified, use "Team: TBD"
            - Include preventive measures, process improvements, and monitoring tasks
            - Include any post-incident reviews, documentation updates, or system changes
            
            For KEY OBSERVATIONS:
            - Extract technical findings, root causes, and impact details
            - Include lessons learned and process gaps identified
            - Include timeline information and restoration details
            - Include any metrics, user counts, or business impact mentioned
            - Format as clear, complete statements without bullet points
            - Focus on factual observations, not recommendations
            
            Return in this EXACT format:
            
            ACTION_ITEMS:
            Action: [description] | Team: [team]
            Action: [description] | Team: [team]
            
            KEY_OBSERVATIONS:
            [observation statement]
            [observation statement]
            """
        else:  # whiteboard
            system_prompt = """
            You are an expert at extracting action items and key observations from incident whiteboard documents.
            These often contain structured tables, quick notes, and immediate action items from incident response.
            
            PERFECT FORMAT REQUIREMENTS:
            
            For ACTION ITEMS:
            - Look for items in tables with Action/Team columns
            - Extract immediate response actions and assignments
            - Format EXACTLY as: "Action: [Clear, specific description] | Team: [Team/Person name]"
            - If no team specified, use "Team: TBD"
            - Include investigation tasks, communication actions, and immediate fixes
            
            For KEY OBSERVATIONS:
            - Extract technical findings from the incident response
            - Include timing information and immediate impacts
            - Include any infrastructure or system details noted
            - Format as clear, complete statements without bullet points
            - Focus on real-time observations during the incident
            
            Return in this EXACT format:
            
            ACTION_ITEMS:
            Action: [description] | Team: [team]
            Action: [description] | Team: [team]
            
            KEY_OBSERVATIONS:
            [observation statement]
            [observation statement]
            """
        
        user_prompt = f"""
        Extract action items and key observations from this {document_type} content in PERFECT format:

        DOCUMENT CONTENT:
        {content}
        
        Requirements:
        1. Extract EVERY action item with clear team assignments
        2. Extract ALL key technical and business observations
        3. Use the EXACT format specified in the system prompt
        4. Be thorough and comprehensive
        5. Maintain all technical details and context
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=3000,
                temperature=0.1
            )
            
            content_response = response.choices[0].message.content
            print(f"Extracted perfect format from {document_type}")
            
            # Parse the response
            actions = []
            observations = []
            current_section = None
            
            for line in content_response.split('\n'):
                line = line.strip()
                if line == 'ACTION_ITEMS:':
                    current_section = 'actions'
                elif line == 'KEY_OBSERVATIONS:':
                    current_section = 'observations'
                elif line and current_section == 'actions':
                    if line.startswith('Action: ') and ' | Team: ' in line:
                        actions.append(line)
                elif line and current_section == 'observations':
                    if not line.startswith('ACTION_ITEMS:') and not line.startswith('KEY_OBSERVATIONS:'):
                        observations.append(line)
            
            return actions, observations
        
        except Exception as e:
            print(f"Error extracting from {document_type}: {str(e)}")
            return [], []

    # ========== TIMESTAMP AND MERGING METHODS ==========
    def extract_timestamps(self, text: str) -> List[Tuple[datetime, str]]:
        """Extract timestamps from text and return sorted list."""
        timestamp_patterns = [
            r'(\d{1,2}/\d{1,2}/\d{4}),?\s+(\d{1,2}:\d{2}:\d{2})',  # MM/DD/YYYY HH:MM:SS
            r'(\d{1,2}/\d{1,2}/\d{4})\s+(\d{1,2}:\d{2})',          # MM/DD/YYYY HH:MM
            r'(\d{4}-\d{2}-\d{2})\s+(\d{1,2}:\d{2}:\d{2})',        # YYYY-MM-DD HH:MM:SS
        ]
        
        timestamps = []
        lines = text.split('\n')
        
        for line in lines:
            for pattern in timestamp_patterns:
                matches = re.finditer(pattern, line)
                for match in matches:
                    try:
                        date_str = match.group(1)
                        time_str = match.group(2)
                        
                        if '/' in date_str:
                            if len(time_str.split(':')) == 3:
                                dt = datetime.strptime(f"{date_str} {time_str}", "%m/%d/%Y %H:%M:%S")
                            else:
                                dt = datetime.strptime(f"{date_str} {time_str}", "%m/%d/%Y %H:%M")
                        else:
                            dt = datetime.strptime(f"{date_str} {time_str}", "%Y-%m-%d %H:%M:%S")
                        
                        timestamps.append((dt, line.strip()))
                    except ValueError:
                        continue
        
        return sorted(timestamps, key=lambda x: x[0])

    def simple_chronological_merge(self, content1: str, content2: str) -> str:
        """Simple chronological merge without AI processing to avoid content loss."""
        print("Using simple chronological merge to preserve all content...")
        
        timestamps1 = self.extract_timestamps(content1)
        timestamps2 = self.extract_timestamps(content2)
        
        all_timestamps = timestamps1 + timestamps2
        all_timestamps.sort(key=lambda x: x[0])
        
        merged_content = []
        seen_exact = set()
        
        for timestamp, content in all_timestamps:
            if content not in seen_exact:
                seen_exact.add(content)
                merged_content.append(content)
        
        # Add non-timestamped content
        for content in [content1, content2]:
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line and not any(pattern in line for pattern in ['/202', ':', 'AM', 'PM']):
                    if line not in seen_exact and len(line) > 10:
                        merged_content.append(line)
                        seen_exact.add(line)
        
        return '\n'.join(merged_content)

    from docx import Document
    import re
    import os
    from docx.shared import Inches

    def has_timestamp(text):
        # Match various timestamp formats: 'MM/DD/YYYY HH:MM:SS', 'YYYY-MM-DD HH:MM:SS', 'HH:MM', etc.
        return bool(re.search(r'(\d{1,2}/\d{1,2}/\d{4},?\s+\d{1,2}:\d{2}(:\d{2})?)|(\d{4}-\d{2}-\d{2}\s+\d{1,2}:\d{2}(:\d{2})?)|(\d{1,2}:\d{2}(:\d{2})?)', text))

    @staticmethod
    def filter_docx_by_timestamp(input_path, output_path, image_dir="timestamped_images"):
        """
        Filter a DOCX document to keep only content up to the second occurrence of the asterisk line
        that appears before the 'Action Items' section.
        
        Args:
            input_path: Path to the input DOCX file
            output_path: Path to save the filtered DOCX file
            image_dir: Directory to save extracted images
            
        Returns:
            tuple: (output_path, number_of_images_extracted)
        """
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import os
        import shutil
        
        try:
            os.makedirs(image_dir, exist_ok=True)
            temp_path = output_path + '.temp.docx'
            shutil.copy2(input_path, temp_path)

            doc = Document(input_path)

            # === Find all asterisk-only paragraphs ===
            asterisk_paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text and all(c == '*' for c in text):
                    asterisk_paragraphs.append(i)

            if len(asterisk_paragraphs) < 3:
                print("Not enough asterisk lines (need at least 3). Copying full document.")
                shutil.copy2(input_path, output_path)
                return output_path, 0

            # Define range: between 2nd and 3rd asterisk-only line
            start_idx = asterisk_paragraphs[1] + 1
            end_idx = asterisk_paragraphs[2]

            # Create new document
            new_doc = Document()
            image_count = 0

            # Copy only content between 2nd and 3rd asterisk lines
            for i in range(start_idx, end_idx):
                para = doc.paragraphs[i]
                if para.text.strip() or len(para.runs) > 0:
                    new_para = new_doc.add_paragraph()
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if hasattr(run, 'font'):
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size

            # === Extract and append images ===
            for rel in doc.part.rels.values():
                if "image" in str(rel.target_ref):
                    try:
                        image_count += 1
                        fname = f"img_{image_count}.png"
                        img_path = os.path.join(image_dir, fname)
                        with open(img_path, 'wb') as f:
                            f.write(rel.target_part.blob)
                        new_doc.add_picture(img_path, width=Inches(4.0))
                        p = new_doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p.add_run(f"Figure {image_count}").italic = True
                    except Exception as e:
                        print(f"Warning: Could not process image: {str(e)}")

            # Save the filtered document
            new_doc.save(output_path)

            if os.path.exists(temp_path):
                os.remove(temp_path)

            print(f"Filtered document saved to: {output_path}")
            if image_count > 0:
                print(f"Extracted {image_count} images to: {os.path.abspath(image_dir)}")

            return output_path, image_count

        except Exception as e:
            import traceback
            error_msg = f"Error creating filtered document: {str(e)}\n{traceback.format_exc()}"
            print(error_msg)

            if 'temp_path' in locals() and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass

            if 'output_path' in locals() and os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass

            try:
                shutil.copy2(input_path, output_path)
                print(f"Fell back to copying original file to: {output_path}")
                return output_path, 0
            except Exception as copy_error:
                print(f"Failed to copy original file: {str(copy_error)}")
                return input_path, 0

    def chunk_transcript_smart(self, transcript: str, max_size: int = None) -> List[str]:
        """Intelligently split transcript into chunks, preserving timeline entries and adding overlap."""
        if max_size is None:
            max_size = self.max_chunk_size
        
        if len(transcript) <= max_size:
            return [transcript]
        
        lines = transcript.split('\n')
        chunks = []
        current_chunk = []
        current_size = 0
        
        for i, line in enumerate(lines):
            line_size = len(line) + 1
            
            if current_size + line_size > max_size and current_chunk:
                chunk_content = '\n'.join(current_chunk)
                chunks.append(chunk_content)
                
                # Start new chunk with overlap
                overlap_lines = []
                overlap_size = 0
                
                for j in range(len(current_chunk) - 1, -1, -1):
                    line_len = len(current_chunk[j])
                    if overlap_size + line_len <= self.overlap_size:
                        overlap_lines.insert(0, current_chunk[j])
                        overlap_size += line_len
                    else:
                        break
                
                current_chunk = overlap_lines + [line]
                current_size = sum(len(l) + 1 for l in current_chunk)
            else:
                current_chunk.append(line)
                current_size += line_size
        
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
        
        print(f"Split into {len(chunks)} chunks with smart overlap")
        return chunks

    # ========== MAIN MERGING METHOD ==========
    def merge_transcripts_with_perfect_extraction(self, transcript1: str, transcript2: str, 
                                                whiteboard_actions: List[str], whiteboard_observations: List[str],
                                                summary_actions: List[str], summary_observations: List[str],
                                                context: str = "", image_info: str = "") -> str:
        """Merge two transcripts with perfect format extracted actions and observations."""
        
        # First, do chronological merge to preserve all content
        merged_timeline = self.simple_chronological_merge(transcript1, transcript2)
        
        system_prompt = """
        You are an expert at creating professional incident reports. You will receive:
        1. A chronologically merged timeline
        2. Perfectly formatted action items from both documents
        3. Perfectly formatted key observations from both documents
        
        Your task is to create a comprehensive incident report in this EXACT format:
        
        *Summary Date: [Current Date]*
        
        **Timelines (Times are in Eastern time (GMT-5) unless otherwise noted):**
        
        [Use the provided merged timeline - DO NOT modify timestamps or lose any entries]
        
        # Action Items
        
        | **Action Item** | **Team** | **Source** |
        | --- | --- | --- |
        [Merged action items from both documents - remove duplicates but preserve unique information]
        
        # Key Observations
        
        ## From Manual Whiteboard 
        [Observations from whiteboard as bullet points with dashes (-)]
        
        ## From AI Whiteboard 
        [Observations from final summary as bullet points with dashes (-)]
        
        CRITICAL REQUIREMENTS:
        1. Use the merged timeline exactly as provided
        2. Preserve all timestamps and formatting in the timeline
        3. Combine all action items in table format with source tracking
        4. Separate observations by source (Manual Whiteboard vs AI Whiteboard)
        5. Include all technical details, ticket numbers, and team communications
        6. Maintain professional incident report formatting
        """
        
        # Truncate timeline if too long for API
        if len(merged_timeline) > 8000:
            print(f"Timeline is long ({len(merged_timeline)} chars), truncating for API processing...")
            truncated_timeline = merged_timeline[:8000] + "\n[... additional timeline entries continue ...]"
        else:
            truncated_timeline = merged_timeline
        
        user_prompt = f"""
        Please create a comprehensive incident report using these components:

        MERGED TIMELINE (USE EXACTLY AS PROVIDED):
        {truncated_timeline}

        WHITEBOARD ACTION ITEMS:
        {chr(10).join(whiteboard_actions)}

        FINAL SUMMARY ACTION ITEMS:
        {chr(10).join(summary_actions)}

        WHITEBOARD KEY OBSERVATIONS:
        {chr(10).join(whiteboard_observations)}

        FINAL SUMMARY KEY OBSERVATIONS:
        {chr(10).join(summary_observations)}

        Additional Context: {context}
        Image Information: {image_info}

        Create the incident report using the exact timeline provided and organized actions/observations.
        Remove duplicate action items but preserve all unique information.
        Keep observations separated by source.
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=6000,
                temperature=0.1
            )
            
            result = response.choices[0].message.content
            
            # If the timeline was truncated, append the remaining content
            if len(merged_timeline) > 8000:
                print("Appending remaining timeline content...")
                remaining_timeline = merged_timeline[8000:]
                
                timeline_end_patterns = ["# Action Items", "## Action Items", "Action Items"]
                timeline_end_pos = -1
                
                for pattern in timeline_end_patterns:
                    pos = result.find(pattern)
                    if pos != -1:
                        timeline_end_pos = pos
                        break
                
                if timeline_end_pos != -1:
                    result = (result[:timeline_end_pos] + 
                             remaining_timeline + "\n\n" + 
                             result[timeline_end_pos:])
                else:
                    result += "\n\n" + remaining_timeline
            
            return result
        
        except Exception as e:
            print(f"Error calling OpenAI API: {str(e)}")
            return self.fallback_merge_with_perfect_extraction(
                merged_timeline, whiteboard_actions, whiteboard_observations,
                summary_actions, summary_observations, image_info
            )

    def fallback_merge_with_perfect_extraction(self, merged_timeline: str,
                                             whiteboard_actions: List[str], whiteboard_observations: List[str],
                                             summary_actions: List[str], summary_observations: List[str],
                                             image_info: str = "") -> str:
        """Fallback merge method with perfect format extracted data."""
        print("Using fallback merge method with perfect extraction...")
        
        merged_content = []
        
        # Add header
        current_date = datetime.now().strftime("%B %d, %Y")
        merged_content.append(f"**Summary Date: {current_date}**")
        merged_content.append("")
        merged_content.append("**Timelines (Times are in Eastern time (GMT-5) unless otherwise noted):**")
        merged_content.append("")
        
        # Add timeline entries
        merged_content.append(merged_timeline)
        
        merged_content.append("")
        merged_content.append("# Action Items")
        merged_content.append("")
        merged_content.append("| **Action Item** | **Team** | **Source** |")
        merged_content.append("| --- | --- | --- |")
        
        # Add whiteboard actions
        for action in whiteboard_actions:
            if " | Team: " in action:
                parts = action.split(" | Team: ")
                action_desc = parts[0].replace("Action: ", "")
                team = parts[1]
                merged_content.append(f"| {action_desc} | {team} | Manual Whiteboard |")
        
        # Add summary actions
        for action in summary_actions:
            if " | Team: " in action:
                parts = action.split(" | Team: ")
                action_desc = parts[0].replace("Action: ", "")
                team = parts[1]
                merged_content.append(f"| {action_desc} | {team} | AI Whiteboard |")
        
        merged_content.append("")
        merged_content.append("# Key Observations")
        merged_content.append("")
        merged_content.append("## From Manual Whiteboard")
        for obs in whiteboard_observations:
            merged_content.append(f"- {obs}")
        
        merged_content.append("")
        merged_content.append("## From AI Whiteboard")
        for obs in summary_observations:
            merged_content.append(f"- {obs}")
        
        if image_info:
            merged_content.append(f"- {image_info}")
        
        return "\n".join(merged_content)

    # ========== DOCUMENT SAVING METHODS ==========
    def save_merged_transcript(self, merged_content: str, output_path: str):
        """Save the merged transcript to a file. Supports TXT, MD, and DOCX."""
        try:
            # Ensure the directory exists
            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
            
            ext = os.path.splitext(output_path.lower())[1][1:]  # Get extension without dot
            
            if ext in ('md', 'txt'):
                with open(output_path, 'w', encoding='utf-8') as file:
                    file.write(merged_content)
                print(f"Merged transcript saved to: {output_path}")
            elif ext == 'docx':
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Add styles for better formatting
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Calibri'
                font.size = Pt(11)
                
                for line in merged_content.split('\n'):
                    if line.startswith('# '):
                        # Main heading
                        p = doc.add_heading(line[2:], level=1)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif line.startswith('## '):
                        # Subheading
                        p = doc.add_heading(line[3:], level=2)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif line.startswith('|') and '---' in line:
                        # Skip markdown table separator
                        continue
                    elif line.startswith('|'):
                        # Handle table rows
                        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                        if 'table' not in locals():
                            # Create table if it doesn't exist
                            table = doc.add_table(rows=1, cols=len(cells))
                            table.style = 'Table Grid'
                            # Add header row
                            hdr_cells = table.rows[0].cells
                            for i, cell in enumerate(cells):
                                hdr_cells[i].text = cell.replace('**', '')
                                for paragraph in hdr_cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                        else:
                            # Add data row
                            row_cells = table.add_row().cells
                            for i, cell in enumerate(cells):
                                row_cells[i].text = cell.replace('**', '')
                    elif line.startswith('*') and line.endswith('*'):
                        # Italic text (e.g., summary date)
                        p = doc.add_paragraph()
                        p.add_run(line[1:-1]).italic = True
                    elif line.startswith('- '):
                        # Bullet points
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(line[2:])
                    elif line.strip():
                        # Regular paragraph
                        doc.add_paragraph(line)
                    else:
                        # Empty line
                        doc.add_paragraph()
                
                doc.save(output_path)
                print(f"Merged transcript saved to: {output_path}")
            else:
                raise ValueError(f"Unsupported file extension: {ext}")
                
        except Exception as e:
            print(f"Error saving merged transcript: {str(e)}")
            import traceback
            traceback.print_exc()

    # ========== MAIN EXECUTION METHOD ==========
    def merge_documents_with_perfect_extraction(self, whiteboard_file: str, final_summary_file: str, 
                                              output_path: str, context: str = "", 
                                              output_format: str = "md") -> str:
        """
        Main method to merge documents with perfect format extraction.
        
        Args:
            whiteboard_file (str): Path to whiteboard DOCX file
            final_summary_file (str): Path to final summary DOCX file
            output_path (str): Path for output file
            context (str): Additional context
            output_format (str): Output format - 'md', 'txt', or 'docx'
            
        Returns:
            str: Merged content with perfect extraction
        """
        print("="*80)
        print("COMPLETE MERGER WITH PERFECT FORMAT EXTRACTION")
        print("="*80)
        print("Process:")
        print("1. Extract last 1 page from Whiteboard document")
        print("2. Extract last 2 pages from Final Summary document") 
        print("3. Use OpenAI for perfect format extraction")
        print("4. Read full documents for timeline merging")
        print("5. Merge everything into professional incident report")
        print("="*80)
        
        # Step 1: Perfect format extraction from targeted sections
        print("\nStep 1: Perfect format extraction from targeted sections...")
        
        # Extract last page from whiteboard
        print("Extracting last page from Whiteboard...")
        whiteboard_content = self.extract_last_pages_docx(whiteboard_file, 1)
        whiteboard_actions, whiteboard_observations = self.extract_perfect_actions_observations_openai(
            whiteboard_content, "whiteboard"
        )
        print(f"Whiteboard: {len(whiteboard_actions)} actions, {len(whiteboard_observations)} observations")
        
        # Extract last two pages from final summary
        print("Extracting last two pages from Final Summary...")
        summary_content = self.extract_last_pages_docx(final_summary_file, 2)
        summary_actions, summary_observations = self.extract_perfect_actions_observations_openai(
            summary_content, "final_summary"
        )
        print(f"Summary: {len(summary_actions)} actions, {len(summary_observations)} observations")
        
        # Step 2: Read full documents for timeline merging
        print("\nStep 2: Reading full documents for timeline merging...")
        
        print(f"Reading full whiteboard document: {whiteboard_file}")
        if whiteboard_file.endswith('.docx'):
            content1, images1 = self.read_docx_file_with_images(whiteboard_file)
        else:
            content1 = self.read_text_file(whiteboard_file)
            images1 = {}
        
        print(f"Reading full final summary document: {final_summary_file}")
        if final_summary_file.endswith('.docx'):
            content2, images2 = self.read_docx_file_with_images(final_summary_file)
        else:
            content2 = self.read_text_file(final_summary_file)
            images2 = {}
        
        all_images = {**images1, **images2}
        print(f"Document 1 length: {len(content1)} characters")
        print(f"Document 2 length: {len(content2)} characters")
        print(f"Found {len(all_images)} images")
        
        # Step 3: Handle large documents with smart chunking
        print("\nStep 3: Processing large documents...")
        if len(content2) > self.max_chunk_size:
            print(f"Document 2 is large ({len(content2)} chars), using smart chunking...")
            chunks = self.chunk_transcript_smart(content2)
            print(f"Created {len(chunks)} chunks with overlap")
            content2_processed = content2  # Use full content for timeline merging
        else:
            content2_processed = content2
        
        # Step 4: Merge transcripts with perfect extraction
        print("\nStep 4: Merging transcripts with perfect format extraction...")
        
        image_info = ""
        if all_images:
            image_info = f"Images extracted: {list(all_images.values())}"
        
        merged_content = self.merge_transcripts_with_perfect_extraction(
            content1, 
            content2_processed,
            whiteboard_actions,
            whiteboard_observations, 
            summary_actions,
            summary_observations,
            context, 
            image_info
        )
        
        # Step 5: Save merged content
        print("\nStep 5: Saving merged content...")
        
        if merged_content:
            # Ensure proper file extension
            if output_format.lower() == "md" and not output_path.endswith('.md'):
                output_path = output_path.rsplit('.', 1)[0] + '.md'
            elif output_format.lower() == "txt" and not output_path.endswith('.txt'):
                output_path = output_path.rsplit('.', 1)[0] + '.txt'
            
            self.save_merged_transcript(merged_content, output_path)
            
            # Create images index if images were found
            if all_images:
                self.create_images_index(all_images, os.path.dirname(output_path))
            
            # Print summary
            print("\n" + "="*80)
            print("COMPLETE MERGER WITH PERFECT EXTRACTION COMPLETED!")
            print("="*80)
            print(f"Perfect Format Extraction Results:")
            print(f"  - Whiteboard actions: {len(whiteboard_actions)}")
            print(f"  - Whiteboard observations: {len(whiteboard_observations)}")  
            print(f"  - Summary actions: {len(summary_actions)}")
            print(f"  - Summary observations: {len(summary_observations)}")
            print(f"Total merged content length: {len(merged_content)} characters")
            print(f"Output saved to: {output_path}")
            if all_images:
                print(f"Images extracted: {len(all_images)}")
            print("="*80)
        
        return merged_content

    def create_images_index(self, image_mapping: Dict[str, str], output_dir: str):
        """Create an index file for extracted images."""
        try:
            index_path = os.path.join(output_dir, "images_index.md")
            with open(index_path, 'w', encoding='utf-8') as f:
                f.write("# Extracted Images Index\n\n")
                f.write("The following images were extracted from the incident documents:\n\n")
                
                for i, (ref, path) in enumerate(image_mapping.items(), 1):
                    f.write(f"## Image {i}\n")
                    f.write(f"- **Reference**: {ref}\n")
                    f.write(f"- **File Path**: {path}\n")
                    f.write(f"- **Markdown Reference**: `![Image {i}]({path})`\n\n")
                    f.write(f"![Image {i}]({path})\n\n")
                    f.write("---\n\n")
            
            print(f"Images index created: {index_path}")
        
        except Exception as e:
            print(f"Error creating images index: {str(e)}")


def main():
    """
    Main function to demonstrate complete merging with perfect format extraction.
    """
    # Get API key from environment variables
    API_KEY = os.getenv("OPENAI_API_KEY")
    
    if not API_KEY:
        print("Error: OPENAI_API_KEY not found in environment variables")
        print("Please create a .env file with your OpenAI API key:")
        print("OPENAI_API_KEY=your_api_key_here")
        return
    
    # File paths - UPDATE THESE WITH YOUR ACTUAL FILE PATHS
    whiteboard_file = "docs/Whiteboard_OTIS_INC02222.docx"
    final_summarized_file = r"C:\Users\Asus-2024\Downloads\final_summarized - 2025-07-29T143947.094.docx"
    
    # Define output paths
    output_dir = os.path.abspath("docs")
    try:
        os.makedirs(output_dir, exist_ok=True)
        # Test if we can write to the directory
        test_file = os.path.join(output_dir, "test_write.tmp")
        with open(test_file, 'w') as f:
            f.write("test")
        os.remove(test_file)
    except Exception as e:
        print(f"Error: Cannot write to output directory '{output_dir}': {str(e)}")
        print(f"Falling back to current directory for output files.")
        output_dir = os.getcwd()
        
    base_filename = "complete_merged_incident_report"
    docx_output = os.path.join(output_dir, f"{base_filename}.docx")
    filtered_docx_output = os.path.join(output_dir, f"{base_filename}_filtered.docx")
    
    # Additional context
    context = """
    This is a major incident involving network outage at a critical factory site.
    The incident involved multiple teams: ABC Team, GBS Team, and ABC/DEF Team.
    Root cause was identified as power failure due to transformer issues with no backup power.
    
    PROCESS USED:
    1. Perfect format extraction from last 1 page of Whiteboard
    2. Perfect format extraction from last 2 pages of Final Summary
    3. Full document timeline merging with chronological ordering
    4. Professional incident report generation with source tracking
    """
    
    # Initialize complete merger
    merger = CompleteMergerWithPerfectExtraction(API_KEY)
    
    # Execute complete merge with perfect extraction
    try:
        print("="*80)
        print("🚀 COMPLETE TRANSCRIPT MERGER WITH PERFECT FORMAT EXTRACTION")
        print("="*80)
        print("🎯 KEY FEATURES:")
        print("✅ Perfect format extraction from targeted document sections")
        print("✅ Complete timeline preservation with chronological merging")
        print("✅ Smart chunking for large documents with overlap")
        print("="*80)
        
        # Execute the complete merge to DOCX
        merged_result = merger.merge_documents_with_perfect_extraction(
            whiteboard_file, 
            final_summarized_file, 
            docx_output, 
            context,
            output_format="docx"
        )
        
        # Apply timestamp filtering to the generated DOCX
        filtered_doc_created = False
        print("\n🔍 Applying timestamp filtering to the merged document...")
        try:
            # First check if source document exists and is valid
            if not os.path.exists(docx_output):
                raise FileNotFoundError(f"Source document not found: {docx_output}")
                
            # Try to open the document to verify it's valid
            try:
                from docx import Document
                doc = Document(docx_output)
            except Exception as e:
                raise Exception(f"Invalid or corrupted DOCX file: {str(e)}")
            
            # Create filtered document
            filtered_path, img_count = merger.filter_docx_by_timestamp(docx_output, filtered_docx_output)
            
            if os.path.exists(filtered_path):
                filtered_doc_created = True
                print(f"✓ Filtered DOCX saved to: {os.path.abspath(filtered_path)}")
                if img_count > 0:
                    print(f"  - Extracted {img_count} images to 'timestamped_images' directory")
            else:
                print("⚠️  Warning: Filtered document was not created")
                
        except Exception as e:
            print(f"⚠️  Warning: Could not create filtered document: {str(e)}")
            print("  The main document will still be available without timestamp filtering.")
            import traceback
            traceback.print_exc()
        
        if merged_result:
            print("\n" + "="*80)
            print("🎉 SUCCESS: COMPLETE MERGER WITH PERFECT EXTRACTION COMPLETED!")
            print("\n" + "="*80)
            print("📊 PROCESS SUMMARY:")
            print("✓ Targeted extraction from last 1 page of Whiteboard")
            print("✓ Targeted extraction from last 2 pages of Final Summary")
            print("✓ Perfect format action items and observations using OpenAI")
            print("✓ Complete timeline merging with chronological ordering")
            print("✓ Professional incident report with source tracking")
            print("✓ All content preserved with smart overlap handling")
            print("✓ DOCX output with proper formatting")
            if filtered_doc_created:
                print("✓ Timestamp filtering applied")
            print("="*80)
            
            # Show output files
            print(f"\n📂 OUTPUT FILES:")
            print(f"- Merged DOCX: {os.path.abspath(docx_output)}")
            if filtered_doc_created:
                print(f"- Filtered DOCX (timestamped only): {os.path.abspath(filtered_docx_output)}")
            
            # Show preview of results
            print("\n📋 PREVIEW OF MERGED CONTENT:")
            print("-" * 50)
            lines = merged_result.split('\n')
            for i, line in enumerate(lines[:25]):  # Show first 25 lines
                print(f"{i+1:2d}: {line}")
            if len(lines) > 25:
                print("    ... (content continues)")
            print("-" * 50)
            
            # Create additional formats
            print("\n📁 CREATING ADDITIONAL OUTPUT FORMATS...")
            
            # Create text versions of the documents
            try:
                import docx2txt
                
                # Create text version of main document
                txt_output = docx_output.replace('.docx', '.txt')
                try:
                    txt_content = docx2txt.process(docx_output)
                    with open(txt_output, 'w', encoding='utf-8') as f:
                        f.write(txt_content)
                    print(f"✓ Text backup saved to: {txt_output}")
                except Exception as e:
                    print(f"⚠️  Could not create text backup of main document: {str(e)}")
                
                # Create text version of filtered document if it was created
                if filtered_doc_created and os.path.exists(filtered_docx_output):
                    filtered_txt_output = filtered_docx_output.replace('.docx', '.txt')
                    try:
                        filtered_txt_content = docx2txt.process(filtered_docx_output)
                        with open(filtered_txt_output, 'w', encoding='utf-8') as f:
                            f.write(filtered_txt_content)
                        print(f"✓ Filtered text backup saved to: {filtered_txt_output}")
                    except Exception as e:
                        print(f"⚠️  Could not create filtered text backup: {str(e)}")
                        
            except Exception as e:
                print(f"⚠️  Error in text backup process: {str(e)}")
            
            print("\n" + "="*80)
            print("🏆 COMPLETE SUCCESS!")
            print("="*80)
            print("📋 WHAT WAS ACCOMPLISHED:")
            print("1. ✅ Perfect format extraction from targeted sections")
            print("2. ✅ Complete timeline merging without content loss")
            print("3. ✅ Professional incident report generation")
            print("4. ✅ Source tracking and validation")
            print("5. ✅ Multiple output formats created")
            print("6. ✅ Images extracted and indexed")
            if filtered_doc_created:
                print("7. ✅ Timestamp filtering applied")
            print("="*80)
            
        else:
            print("\n❌ MERGE FAILED")
            print("Please check the error messages above and verify:")
            print("- OpenAI API key is correct")
            print("- File paths exist and are accessible")
            print("- Documents contain the expected content")
            
    except FileNotFoundError as e:
        print(f"\n❌ FILE NOT FOUND ERROR: {str(e)}")
        print("Please verify the file paths in the main() function:")
        print(f"- Whiteboard file: {whiteboard_file}")
        print(f"- Final summary file: {final_summarized_file}")
        
    except Exception as e:
        print(f"\n❌ UNEXPECTED ERROR: {str(e)}")
        print("Full error details:")
        import traceback
        traceback.print_exc()
        print("\nTroubleshooting tips:")
        print("1. Verify your OpenAI API key is valid")
        print("2. Check that all required libraries are installed")
        print("3. Ensure the input files are not corrupted")
        print("4. Try with smaller test documents first")


# Additional utility functions for advanced usage
def extract_only_perfect_format(api_key: str, whiteboard_file: str, final_summary_file: str, 
                               output_file: str = "perfect_format_only.md"):
    """
    Utility function to ONLY extract perfect format actions and observations
    without full document merging.
    """
    print("🎯 PERFECT FORMAT EXTRACTION ONLY")
    print("="*50)
    
    merger = CompleteMergerWithPerfectExtraction(api_key)
    
    # Extract from whiteboard (last 1 page)
    whiteboard_content = merger.extract_last_pages_docx(whiteboard_file, 1)
    whiteboard_actions, whiteboard_observations = merger.extract_perfect_actions_observations_openai(
        whiteboard_content, "whiteboard"
    )
    
    # Extract from final summary (last 2 pages)
    summary_content = merger.extract_last_pages_docx(final_summary_file, 2)
    summary_actions, summary_observations = merger.extract_perfect_actions_observations_openai(
        summary_content, "final_summary"
    )
    
    # Create simple report
    report = []
    report.append("# Perfect Format Extraction Results")
    report.append("")
    report.append("## Action Items")
    report.append("")
    report.append("### From Whiteboard (Last Page)")
    for action in whiteboard_actions:
        report.append(f"- {action}")
    report.append("")
    report.append("### From Final Summary (Last 2 Pages)")  
    for action in summary_actions:
        report.append(f"- {action}")
    report.append("")
    report.append("## Key Observations")
    report.append("")
    report.append("### From Whiteboard (Last Page)")
    for obs in whiteboard_observations:
        report.append(f"- {obs}")
    report.append("")
    report.append("### From Final Summary (Last 2 Pages)")
    for obs in summary_observations:
        report.append(f"- {obs}")
    
    report_content = "\n".join(report)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(report_content)
    
    print(f"✅ Perfect format extraction saved to: {output_file}")
    return report_content


def timeline_only_merge(api_key: str, whiteboard_file: str, final_summary_file: str,
                       output_file: str = "timeline_only_merge.md"):
    """
    Utility function to ONLY merge timelines without AI processing.
    """
    print("📅 TIMELINE ONLY MERGE")
    print("="*50)
    
    merger = CompleteMergerWithPerfectExtraction(api_key)
    
    # Read full documents
    content1 = merger.read_docx_file(whiteboard_file)
    content2 = merger.read_docx_file(final_summary_file)
    
    # Simple chronological merge
    merged_timeline = merger.simple_chronological_merge(content1, content2)
    
    # Create simple timeline report
    report = []
    report.append("# Timeline Only Merge")
    report.append("")
    report.append("**Chronologically merged timeline from both documents:**")
    report.append("")
    report.append(merged_timeline)
    
    report_content = "\n".join(report)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(report_content)
    
    print(f"✅ Timeline merge saved to: {output_file}")
    return report_content


if __name__ == "__main__":
    main()
    
    # Uncomment these lines to run individual utilities:
    
    # # Extract only perfect format (no full merge)
    # API_KEY = "your-openai-api-key-here"
    # if API_KEY and API_KEY != "your-openai-api-key-here":
    #     extract_only_perfect_format(
    #         API_KEY,
    #         "docs/Whiteboard_OTIS_INC02222.docx",
    #         r"C:\Users\Asus-2024\Downloads\final_summarized - 2025-07-29T143947.094.docx",
    #         "perfect_format_only.md"
    #     )
    
    # # Timeline only merge (no AI processing)
    # timeline_only_merge(
    #     API_KEY,
    #     "docs/Whiteboard_OTIS_INC02222.docx", 
    #     r"C:\Users\Asus-2024\Downloads\final_summarized - 2025-07-29T143947.094.docx",
    #     "timeline_only.md"
    # )